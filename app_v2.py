"""
水质检测报告系统 V2 - 主应用
支持模板管理、权限系统、多格式导出等功能
"""
from flask import Flask, render_template, request, jsonify, send_file, session
from models_v2 import get_db_connection, init_database, DATABASE_PATH
from auth import (
    login_user, logout_user, get_current_user, login_required, admin_required,
    create_user, change_password, log_operation, get_operation_logs
)
from datetime import datetime, timedelta
import json
import os
import shutil
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-in-production'  # 生产环境需修改
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)  # Session有效期7天

# 初始化数据库
init_database()

# ==================== 认证相关 API ====================
@app.route('/api/auth/login', methods=['POST'])
def api_login():
    """用户登录"""
    data = request.json
    username = data.get('username')
    password = data.get('password')

    if not username or not password:
        return jsonify({'error': '用户名和密码不能为空'}), 400

    success, message, user = login_user(username, password)

    if success:
        log_operation('用户登录', f'用户 {username} 登录成功')
        return jsonify({'message': message, 'user': user})
    else:
        return jsonify({'error': message}), 401

@app.route('/api/auth/logout', methods=['POST'])
def api_logout():
    """用户登出"""
    username = session.get('username', '未知用户')
    success, message = logout_user()
    log_operation('用户登出', f'用户 {username} 退出登录', user_id=None)
    return jsonify({'message': message})

@app.route('/api/auth/current-user', methods=['GET'])
def api_current_user():
    """获取当前登录用户"""
    user = get_current_user()
    if user:
        return jsonify({'user': user})
    else:
        return jsonify({'user': None}), 401

@app.route('/api/auth/change-password', methods=['POST'])
@login_required
def api_change_password():
    """修改密码"""
    data = request.json
    old_password = data.get('old_password')
    new_password = data.get('new_password')

    if not old_password or not new_password:
        return jsonify({'error': '旧密码和新密码不能为空'}), 400

    user_id = session['user_id']
    success, message = change_password(user_id, old_password, new_password)

    if success:
        log_operation('修改密码', '用户修改密码成功')
        return jsonify({'message': message})
    else:
        return jsonify({'error': message}), 400

@app.route('/api/users', methods=['GET', 'POST'])
@admin_required
def api_users():
    """用户管理(仅管理员)"""
    if request.method == 'POST':
        data = request.json
        username = data.get('username')
        password = data.get('password')
        role = data.get('role', 'reporter')

        if not username or not password:
            return jsonify({'error': '用户名和密码不能为空'}), 400

        success, message, user_id = create_user(username, password, role)

        if success:
            log_operation('创建用户', f'创建用户 {username}, 角色:{role}')
            return jsonify({'message': message, 'user_id': user_id}), 201
        else:
            return jsonify({'error': message}), 400

    # GET请求 - 获取所有用户
    conn = get_db_connection()
    users = conn.execute('SELECT id, username, role, created_at FROM users').fetchall()
    conn.close()

    return jsonify([dict(user) for user in users])

# ==================== 公司管理 API ====================
@app.route('/api/companies', methods=['GET', 'POST'])
@login_required
def api_companies():
    """公司管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        data = request.json
        name = data.get('name')

        if not name:
            return jsonify({'error': '公司名称不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute('INSERT INTO companies (name) VALUES (?)', (name,))
            conn.commit()
            company_id = cursor.lastrowid

            log_operation('添加公司', f'添加公司: {name}', conn=conn)
            conn.close()

            return jsonify({'id': company_id, 'message': '公司添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': '公司名称已存在'}), 400

    # GET请求
    companies = conn.execute('SELECT * FROM companies ORDER BY name').fetchall()
    conn.close()

    return jsonify([dict(company) for company in companies])

@app.route('/api/companies/<int:id>', methods=['PUT', 'DELETE'])
@admin_required
def api_company_detail(id):
    """公司详情操作"""
    conn = get_db_connection()

    if request.method == 'DELETE':
        company = conn.execute('SELECT name FROM companies WHERE id = ?', (id,)).fetchone()

        if not company:
            conn.close()
            return jsonify({'error': '公司不存在'}), 404

        conn.execute('DELETE FROM companies WHERE id = ?', (id,))
        conn.commit()

        log_operation('删除公司', f'删除公司: {company["name"]}', conn=conn)
        conn.close()

        return jsonify({'message': '公司删除成功'})

    if request.method == 'PUT':
        data = request.json
        name = data.get('name')

        if not name:
            conn.close()
            return jsonify({'error': '公司名称不能为空'}), 400

        try:
            conn.execute('UPDATE companies SET name = ? WHERE id = ?', (name, id))
            conn.commit()
            conn.close()

            log_operation('更新公司', f'更新公司: {name}')
            return jsonify({'message': '公司更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': '公司名称已存在'}), 400

# ==================== 客户管理 API ====================
@app.route('/api/customers', methods=['GET', 'POST'])
@login_required
def api_customers():
    """客户管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        data = request.json
        inspected_unit = data.get('inspected_unit', '').strip()
        water_plant = data.get('water_plant', '').strip()
        unit_address = data.get('unit_address', '').strip()
        contact_person = data.get('contact_person', '').strip()
        contact_phone = data.get('contact_phone', '').strip()
        email = data.get('email', '').strip()
        remark = data.get('remark', '').strip()

        if not inspected_unit:
            conn.close()
            return jsonify({'error': '被检单位不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO customers (inspected_unit, water_plant, unit_address,
                                      contact_person, contact_phone, email, remark)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (inspected_unit, water_plant, unit_address, contact_person,
                  contact_phone, email, remark))
            conn.commit()
            customer_id = cursor.lastrowid

            log_operation('添加客户', f'添加客户: {inspected_unit}', conn=conn)
            conn.close()

            return jsonify({'id': customer_id, 'message': '客户添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': f'添加客户失败: {str(e)}'}), 400

    # GET请求
    customers = conn.execute('''
        SELECT id, inspected_unit, water_plant, unit_address,
               contact_person, contact_phone, email, remark,
               created_at, updated_at
        FROM customers
        ORDER BY created_at DESC
    ''').fetchall()
    conn.close()

    return jsonify([dict(customer) for customer in customers])

@app.route('/api/customers/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_customer_detail(id):
    """客户详情操作"""
    conn = get_db_connection()

    # GET请求 - 获取单个客户详情
    if request.method == 'GET':
        customer = conn.execute('''
            SELECT id, inspected_unit, water_plant, unit_address,
                   contact_person, contact_phone, email, remark,
                   created_at, updated_at
            FROM customers WHERE id = ?
        ''', (id,)).fetchone()

        if not customer:
            conn.close()
            return jsonify({'error': '客户不存在'}), 404

        conn.close()
        return jsonify(dict(customer))

    # DELETE请求
    if request.method == 'DELETE':
        customer = conn.execute('SELECT inspected_unit FROM customers WHERE id = ?', (id,)).fetchone()

        if not customer:
            conn.close()
            return jsonify({'error': '客户不存在'}), 404

        conn.execute('DELETE FROM customers WHERE id = ?', (id,))
        conn.commit()

        log_operation('删除客户', f'删除客户: {customer["inspected_unit"]}', conn=conn)
        conn.close()

        return jsonify({'message': '客户删除成功'})

    # PUT请求
    if request.method == 'PUT':
        data = request.json
        inspected_unit = data.get('inspected_unit', '').strip()
        water_plant = data.get('water_plant', '').strip()
        unit_address = data.get('unit_address', '').strip()
        contact_person = data.get('contact_person', '').strip()
        contact_phone = data.get('contact_phone', '').strip()
        email = data.get('email', '').strip()
        remark = data.get('remark', '').strip()

        if not inspected_unit:
            conn.close()
            return jsonify({'error': '被检单位不能为空'}), 400

        try:
            conn.execute('''
                UPDATE customers
                SET inspected_unit = ?, water_plant = ?, unit_address = ?,
                    contact_person = ?, contact_phone = ?, email = ?, remark = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (inspected_unit, water_plant, unit_address, contact_person,
                  contact_phone, email, remark, id))
            conn.commit()

            log_operation('更新客户', f'更新客户: {inspected_unit}', conn=conn)
            conn.close()

            return jsonify({'message': '客户更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': f'更新客户失败: {str(e)}'}), 400

# ==================== 样品类型管理 API ====================
@app.route('/api/sample-types', methods=['GET', 'POST'])
@login_required
def api_sample_types():
    """样品类型管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        # 仅管理员可创建
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        name = data.get('name')
        code = data.get('code')
        description = data.get('description', '')
        remark = data.get('remark', '')
        indicator_ids = data.get('indicator_ids', [])  # 检测项目ID列表

        if not name or not code:
            return jsonify({'error': '样品类型名称和代码不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO sample_types (name, code, description, remark) VALUES (?, ?, ?, ?)',
                (name, code, description, remark)
            )
            sample_type_id = cursor.lastrowid

            # 添加检测项目关联（使用间隔序号）
            if indicator_ids:
                for idx, indicator_id in enumerate(indicator_ids):
                    cursor.execute(
                        'INSERT INTO template_indicators (sample_type_id, indicator_id, sort_order) VALUES (?, ?, ?)',
                        (sample_type_id, indicator_id, idx * 10)  # 使用间隔序号 0, 10, 20, 30...
                    )

            conn.commit()
            conn.close()

            log_operation('添加样品类型', f'添加样品类型: {name} ({code})，关联{len(indicator_ids)}个检测项目')
            return jsonify({'id': sample_type_id, 'message': '样品类型添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': '样品类型名称或代码已存在'}), 400

    # GET请求 - 支持搜索
    search = request.args.get('search', '')

    if search:
        sample_types = conn.execute(
            'SELECT * FROM sample_types WHERE name LIKE ? OR remark LIKE ? ORDER BY created_at DESC',
            (f'%{search}%', f'%{search}%')
        ).fetchall()
    else:
        sample_types = conn.execute('SELECT * FROM sample_types ORDER BY created_at DESC').fetchall()
    conn.close()

    return jsonify([dict(st) for st in sample_types])

@app.route('/api/sample-types/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_sample_type_detail(id):
    """样品类型详情操作"""
    conn = get_db_connection()

    if request.method == 'GET':
        # 获取样品类型基本信息
        sample_type = conn.execute('SELECT * FROM sample_types WHERE id = ?', (id,)).fetchone()

        if not sample_type:
            conn.close()
            return jsonify({'error': '样品类型不存在'}), 404

        # 获取已关联的检测项目ID列表
        indicator_ids = conn.execute(
            'SELECT indicator_id FROM template_indicators WHERE sample_type_id = ? ORDER BY sort_order',
            (id,)
        ).fetchall()

        result = dict(sample_type)
        result['indicator_ids'] = [row['indicator_id'] for row in indicator_ids]

        conn.close()
        return jsonify(result)

    if request.method == 'DELETE':
        # 仅管理员可删除
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403
        sample_type = conn.execute('SELECT name FROM sample_types WHERE id = ?', (id,)).fetchone()

        if not sample_type:
            conn.close()
            return jsonify({'error': '样品类型不存在'}), 404

        conn.execute('DELETE FROM sample_types WHERE id = ?', (id,))
        conn.commit()

        log_operation('删除样品类型', f'删除样品类型: {sample_type["name"]}', conn=conn)
        conn.close()

        return jsonify({'message': '样品类型删除成功'})

    if request.method == 'PUT':
        # 仅管理员可更新
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        name = data.get('name')
        code = data.get('code')
        description = data.get('description', '')
        remark = data.get('remark', '')
        indicator_ids = data.get('indicator_ids', [])  # 检测项目ID列表
        client_version = data.get('version')  # 客户端的版本号

        if not name or not code:
            return jsonify({'error': '样品类型名称和代码不能为空'}), 400

        try:
            cursor = conn.cursor()

            # 获取当前版本号进行乐观锁检查
            current = cursor.execute(
                'SELECT version FROM sample_types WHERE id = ?', (id,)
            ).fetchone()

            if not current:
                conn.close()
                return jsonify({'error': '样品类型不存在'}), 404

            current_version = current['version'] if current['version'] else 1

            # 版本号检查：如果客户端提供了版本号，检查是否匹配
            if client_version is not None and current_version != client_version:
                conn.close()
                return jsonify({
                    'error': '数据已被其他用户修改，请刷新页面后重试',
                    'conflict': True,
                    'current_version': current_version
                }), 409  # 409 Conflict

            # 更新样品类型，版本号+1
            new_version = current_version + 1
            cursor.execute(
                'UPDATE sample_types SET name = ?, code = ?, description = ?, remark = ?, version = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?',
                (name, code, description, remark, new_version, id)
            )

            # 更新检测项目关联：先删除旧关联，再添加新关联（使用间隔序号）
            cursor.execute('DELETE FROM template_indicators WHERE sample_type_id = ?', (id,))

            if indicator_ids:
                for idx, indicator_id in enumerate(indicator_ids):
                    cursor.execute(
                        'INSERT INTO template_indicators (sample_type_id, indicator_id, sort_order) VALUES (?, ?, ?)',
                        (id, indicator_id, idx * 10)  # 使用间隔序号 0, 10, 20, 30...
                    )

            conn.commit()
            conn.close()

            log_operation('更新样品类型', f'更新样品类型: {name} ({code})，关联{len(indicator_ids)}个检测项目 (v{current_version}->v{new_version})')
            return jsonify({
                'message': '样品类型更新成功',
                'version': new_version
            })
        except Exception as e:
            conn.close()
            if 'UNIQUE constraint failed' in str(e):
                return jsonify({'error': '样品类型名称或代码已存在'}), 400
            return jsonify({'error': f'更新失败: {str(e)}'}), 400

# ==================== 检测项目分组管理 API ====================
@app.route('/api/indicator-groups', methods=['GET', 'POST'])
@login_required
def api_indicator_groups():
    """检测项目分组管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        # 仅管理员可创建
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        name = data.get('name')
        sort_order = data.get('sort_order', 0)

        if not name:
            return jsonify({'error': '分组名称不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO indicator_groups (name, sort_order) VALUES (?, ?)',
                (name, sort_order)
            )
            conn.commit()
            group_id = cursor.lastrowid
            conn.close()

            log_operation('添加检测项目分组', f'添加分组: {name}')
            return jsonify({'id': group_id, 'message': '分组添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': '分组名称已存在'}), 400

    # GET请求
    groups = conn.execute('SELECT * FROM indicator_groups ORDER BY sort_order, name').fetchall()
    conn.close()

    return jsonify([dict(group) for group in groups])

@app.route('/api/indicator-groups/<int:id>', methods=['PUT', 'DELETE'])
@admin_required
def api_indicator_group_detail(id):
    """检测项目分组详情操作"""
    conn = get_db_connection()

    if request.method == 'DELETE':
        group = conn.execute('SELECT name, is_system FROM indicator_groups WHERE id = ?', (id,)).fetchone()

        if not group:
            conn.close()
            return jsonify({'error': '分组不存在'}), 404

        # 检查是否为系统分组
        if group['is_system']:
            conn.close()
            return jsonify({'error': '系统分组不能删除'}), 403

        conn.execute('DELETE FROM indicator_groups WHERE id = ?', (id,))
        conn.commit()

        log_operation('删除检测项目分组', f'删除分组: {group["name"]}', conn=conn)
        conn.close()

        return jsonify({'message': '分组删除成功'})

    if request.method == 'PUT':
        data = request.json
        name = data.get('name')
        sort_order = data.get('sort_order', 0)

        if not name:
            return jsonify({'error': '分组名称不能为空'}), 400

        try:
            conn.execute(
                'UPDATE indicator_groups SET name = ?, sort_order = ? WHERE id = ?',
                (name, sort_order, id)
            )
            conn.commit()
            conn.close()

            log_operation('更新检测项目分组', f'更新分组: {name}')
            return jsonify({'message': '分组更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': '分组名称已存在'}), 400

# ==================== 检测指标管理 API ====================
@app.route('/api/indicators', methods=['GET', 'POST'])
@login_required
def api_indicators():
    """检测指标管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        # 仅管理员可创建
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        group_id = data.get('group_id')
        name = data.get('name')
        unit = data.get('unit', '')
        default_value = data.get('default_value', '')
        limit_value = data.get('limit_value', '')
        detection_method = data.get('detection_method', '')
        description = data.get('description', '')
        remark = data.get('remark', '')
        sort_order = data.get('sort_order', 0)

        if not name:
            return jsonify({'error': '指标名称不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO indicators (group_id, name, unit, default_value, limit_value, detection_method, description, remark, sort_order) '
                'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (group_id, name, unit, default_value, limit_value, detection_method, description, remark, sort_order)
            )
            conn.commit()
            indicator_id = cursor.lastrowid

            log_operation('添加检测指标', f'添加指标: {name}', conn=conn)
            conn.close()

            return jsonify({'id': indicator_id, 'message': '指标添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': '指标名称已存在'}), 400

    # GET请求 - 支持按分组筛选
    group_id = request.args.get('group_id')

    if group_id:
        indicators = conn.execute(
            'SELECT * FROM indicators WHERE group_id = ? ORDER BY sort_order, name',
            (group_id,)
        ).fetchall()
    else:
        indicators = conn.execute(
            'SELECT i.*, g.name as group_name FROM indicators i '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'ORDER BY i.sort_order, i.name'
        ).fetchall()

    conn.close()

    return jsonify([dict(indicator) for indicator in indicators])

@app.route('/api/indicators/<int:id>', methods=['PUT', 'DELETE'])
@admin_required
def api_indicator_detail(id):
    """检测指标详情操作"""
    conn = get_db_connection()

    if request.method == 'DELETE':
        try:
            indicator = conn.execute('SELECT name FROM indicators WHERE id = ?', (id,)).fetchone()

            if not indicator:
                conn.close()
                return jsonify({'error': '指标不存在'}), 404

            # 检查是否被模板使用
            template_usage = conn.execute(
                'SELECT COUNT(*) as count FROM template_indicators WHERE indicator_id = ?',
                (id,)
            ).fetchone()

            if template_usage['count'] > 0:
                conn.close()
                return jsonify({'error': f'该指标正在被 {template_usage["count"]} 个模板使用，无法删除'}), 400

            # 检查是否被报告数据使用
            report_usage = conn.execute(
                'SELECT COUNT(*) as count FROM report_data WHERE indicator_id = ?',
                (id,)
            ).fetchone()

            if report_usage['count'] > 0:
                conn.close()
                return jsonify({'error': f'该指标已在 {report_usage["count"]} 份报告中使用，无法删除'}), 400

            # 执行删除
            conn.execute('DELETE FROM indicators WHERE id = ?', (id,))

            log_operation('删除检测指标', f'删除指标: {indicator["name"]}', conn=conn)
            conn.close()

            return jsonify({'message': '指标删除成功'})
        except Exception as e:
            conn.close()
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'删除失败: {str(e)}'}), 500

    if request.method == 'PUT':
        data = request.json
        group_id = data.get('group_id')
        name = data.get('name')
        unit = data.get('unit', '')
        default_value = data.get('default_value', '')
        limit_value = data.get('limit_value', '')
        detection_method = data.get('detection_method', '')
        description = data.get('description', '')
        remark = data.get('remark', '')
        sort_order = data.get('sort_order', 0)

        if not name:
            return jsonify({'error': '指标名称不能为空'}), 400

        try:
            conn.execute(
                'UPDATE indicators SET group_id = ?, name = ?, unit = ?, default_value = ?, limit_value = ?, detection_method = ?, '
                'description = ?, remark = ?, sort_order = ? WHERE id = ?',
                (group_id, name, unit, default_value, limit_value, detection_method, description, remark, sort_order, id)
            )
            conn.commit()

            log_operation('更新检测指标', f'更新指标: {name}', conn=conn)
            conn.close()

            return jsonify({'message': '指标更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': '指标名称已存在'}), 400

# ==================== 模板-检测项目关联 API ====================
@app.route('/api/template-indicators', methods=['GET', 'POST'])
@login_required
def api_template_indicators():
    """模板检测项目关联"""
    conn = get_db_connection()

    if request.method == 'POST':
        # 仅管理员可创建
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        sample_type_id = data.get('sample_type_id')
        indicator_id = data.get('indicator_id')
        is_required = data.get('is_required', False)
        sort_order = data.get('sort_order', 0)

        if not sample_type_id or not indicator_id:
            return jsonify({'error': '样品类型和检测指标不能为空'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO template_indicators (sample_type_id, indicator_id, is_required, sort_order) '
                'VALUES (?, ?, ?, ?)',
                (sample_type_id, indicator_id, is_required, sort_order)
            )
            conn.commit()
            ti_id = cursor.lastrowid

            log_operation('添加模板检测项', f'样品类型ID:{sample_type_id}, 指标ID:{indicator_id}', conn=conn)
            conn.close()

            return jsonify({'id': ti_id, 'message': '检测项目添加成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': '该检测项目已存在于模板中'}), 400

    # GET请求 - 获取指定样品类型的检测项目
    sample_type_id = request.args.get('sample_type_id')

    if sample_type_id:
        template_indicators = conn.execute(
            'SELECT ti.*, i.name as indicator_name, i.unit, i.default_value, i.limit_value, '
            'i.detection_method, i.group_id, g.name as group_name '
            'FROM template_indicators ti '
            'LEFT JOIN indicators i ON ti.indicator_id = i.id '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'WHERE ti.sample_type_id = ? '
            'ORDER BY ti.sort_order, g.sort_order, i.sort_order',
            (sample_type_id,)
        ).fetchall()
    else:
        template_indicators = conn.execute(
            'SELECT ti.*, i.name as indicator_name, st.name as sample_type_name '
            'FROM template_indicators ti '
            'LEFT JOIN indicators i ON ti.indicator_id = i.id '
            'LEFT JOIN sample_types st ON ti.sample_type_id = st.id '
            'ORDER BY ti.sample_type_id, ti.sort_order'
        ).fetchall()

    conn.close()

    return jsonify([dict(ti) for ti in template_indicators])

@app.route('/api/template-indicators/<int:id>', methods=['DELETE'])
@admin_required
def api_template_indicator_delete(id):
    """删除模板检测项目"""
    conn = get_db_connection()

    conn.execute('DELETE FROM template_indicators WHERE id = ?', (id,))
    conn.commit()

    log_operation('删除模板检测项', f'模板检测项ID:{id}', conn=conn)
    conn.close()

    return jsonify({'message': '检测项目删除成功'})

# ==================== 报告管理 API ====================
@app.route('/api/reports', methods=['GET', 'POST'])
@login_required
def api_reports():
    """报告管理"""
    conn = get_db_connection()

    if request.method == 'POST':
        data = request.json
        sample_number = data.get('sample_number')
        company_id = data.get('company_id')
        sample_type_id = data.get('sample_type_id')
        detection_person = data.get('detection_person', '')
        review_person = data.get('review_person', '')
        detection_date = data.get('detection_date')
        remark = data.get('remark', '')
        report_data_list = data.get('data', [])
        template_id = data.get('template_id')
        template_fields = data.get('template_fields', [])
        review_status = data.get('review_status', 'draft')  # 默认为草稿，可以是 'draft' 或 'pending'

        # 新增字段
        report_date = data.get('report_date')
        sample_source = data.get('sample_source', '')
        sampler = data.get('sampler', '')
        sampling_date = data.get('sampling_date')
        sampling_basis = data.get('sampling_basis', '')
        sample_received_date = data.get('sample_received_date')
        sampling_location = data.get('sampling_location', '')
        sample_status = data.get('sample_status', '')
        product_standard = data.get('product_standard', '')
        test_conclusion = data.get('test_conclusion', '')
        detection_items_description = data.get('detection_items_description', '')
        additional_info = data.get('additional_info', '')

        # 获取用户输入的报告编号
        report_number = data.get('report_number', '').strip()

        if not report_number or not sample_number or not sample_type_id:
            return jsonify({'error': '报告编号、样品编号和样品类型不能为空'}), 400

        # 检查报告编号是否已存在
        existing = conn.execute(
            'SELECT id FROM reports WHERE report_number = ?',
            (report_number,)
        ).fetchone()

        if existing:
            conn.close()
            return jsonify({'error': f'报告编号 {report_number} 已存在'}), 400

        try:
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO reports (report_number, sample_number, company_id, sample_type_id, '
                'detection_person, review_person, detection_date, remark, template_id, review_status, created_by, '
                'report_date, sample_source, sampler, sampling_date, sampling_basis, '
                'sample_received_date, sampling_location, sample_status, product_standard, '
                'test_conclusion, detection_items_description, additional_info) '
                'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (report_number, sample_number, company_id, sample_type_id, detection_person,
                 review_person, detection_date, remark, template_id, review_status, session['user_id'],
                 report_date, sample_source, sampler, sampling_date, sampling_basis,
                 sample_received_date, sampling_location, sample_status, product_standard,
                 test_conclusion, detection_items_description, additional_info)
            )
            conn.commit()
            report_id = cursor.lastrowid

            # 添加报告数据
            for item in report_data_list:
                if item.get('indicator_id'):
                    cursor.execute(
                        'INSERT INTO report_data (report_id, indicator_id, measured_value, remark) '
                        'VALUES (?, ?, ?, ?)',
                        (report_id, item['indicator_id'], item.get('measured_value', ''),
                         item.get('remark', ''))
                    )

            # 添加模板字段值
            for field in template_fields:
                if field.get('field_mapping_id') and field.get('field_value'):
                    cursor.execute(
                        'INSERT INTO report_field_values (report_id, field_mapping_id, field_value) '
                        'VALUES (?, ?, ?)',
                        (report_id, field['field_mapping_id'], field['field_value'])
                    )

            conn.commit()
            conn.close()

            status_text = '草稿' if review_status == 'draft' else '提交审核'
            log_operation('创建报告', f'报告编号:{report_number}, 状态:{status_text}')
            return jsonify({'id': report_id, 'report_number': report_number, 'message': '报告创建成功'}), 201
        except Exception as e:
            conn.close()
            return jsonify({'error': str(e)}), 500

    # GET请求 - 支持搜索
    search_sample_number = request.args.get('sample_number', '')
    search_company_id = request.args.get('company_id', '')
    limit = int(request.args.get('limit', 100))
    offset = int(request.args.get('offset', 0))

    query = '''
        SELECT r.*, st.name as sample_type_name, c.name as company_name
        FROM reports r
        LEFT JOIN sample_types st ON r.sample_type_id = st.id
        LEFT JOIN companies c ON r.company_id = c.id
        WHERE 1=1
    '''
    params = []

    if search_sample_number:
        query += ' AND r.sample_number LIKE ?'
        params.append(f'%{search_sample_number}%')

    if search_company_id:
        query += ' AND r.company_id = ?'
        params.append(search_company_id)

    query += ' ORDER BY r.created_at DESC LIMIT ? OFFSET ?'
    params.extend([limit, offset])

    reports = conn.execute(query, params).fetchall()
    conn.close()

    return jsonify([dict(report) for report in reports])

@app.route('/api/reports/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_report_detail(id):
    """报告详情"""
    conn = get_db_connection()

    if request.method == 'DELETE':
        # 仅创建人或管理员可删除
        report = conn.execute('SELECT created_by, generated_report_path FROM reports WHERE id = ?', (id,)).fetchone()

        if not report:
            conn.close()
            return jsonify({'error': '报告不存在'}), 404

        if session.get('role') != 'admin' and report['created_by'] != session['user_id']:
            conn.close()
            return jsonify({'error': '无权删除此报告'}), 403

        # 删除生成的报告文件（如果存在）
        if report['generated_report_path'] and os.path.exists(report['generated_report_path']):
            try:
                os.remove(report['generated_report_path'])
            except Exception as e:
                print(f"删除报告文件失败: {e}")

        conn.execute('DELETE FROM reports WHERE id = ?', (id,))
        conn.commit()

        log_operation('删除报告', f'报告ID:{id}', conn=conn)
        conn.close()

        return jsonify({'message': '报告删除成功'})

    if request.method == 'PUT':
        # 仅创建人或管理员可修改
        report = conn.execute('SELECT created_by, report_number FROM reports WHERE id = ?', (id,)).fetchone()

        if not report:
            conn.close()
            return jsonify({'error': '报告不存在'}), 404

        if session.get('role') != 'admin' and report['created_by'] != session['user_id']:
            conn.close()
            return jsonify({'error': '无权修改此报告'}), 403

        data = request.json
        sample_number = data.get('sample_number')
        company_id = data.get('company_id')
        detection_person = data.get('detection_person', '')
        review_person = data.get('review_person', '')
        detection_date = data.get('detection_date')
        remark = data.get('remark', '')
        report_data_list = data.get('data', [])
        template_fields = data.get('template_fields', [])

        # 新增字段
        report_date = data.get('report_date')
        sample_source = data.get('sample_source', '')
        sampler = data.get('sampler', '')
        sampling_date = data.get('sampling_date')
        sampling_basis = data.get('sampling_basis', '')
        sample_received_date = data.get('sample_received_date')
        sampling_location = data.get('sampling_location', '')
        sample_status = data.get('sample_status', '')
        product_standard = data.get('product_standard', '')
        test_conclusion = data.get('test_conclusion', '')
        detection_items_description = data.get('detection_items_description', '')
        additional_info = data.get('additional_info', '')

        try:
            cursor = conn.cursor()
            # 更新报告基本信息
            cursor.execute(
                'UPDATE reports SET sample_number = ?, company_id = ?, detection_person = ?, review_person = ?, '
                'detection_date = ?, remark = ?, report_date = ?, sample_source = ?, sampler = ?, '
                'sampling_date = ?, sampling_basis = ?, sample_received_date = ?, sampling_location = ?, '
                'sample_status = ?, product_standard = ?, test_conclusion = ?, detection_items_description = ?, '
                'additional_info = ? WHERE id = ?',
                (sample_number, company_id, detection_person, review_person, detection_date, remark,
                 report_date, sample_source, sampler, sampling_date, sampling_basis,
                 sample_received_date, sampling_location, sample_status, product_standard,
                 test_conclusion, detection_items_description, additional_info, id)
            )

            # 删除旧的报告数据
            cursor.execute('DELETE FROM report_data WHERE report_id = ?', (id,))

            # 插入新的报告数据
            for item in report_data_list:
                if item.get('indicator_id'):
                    cursor.execute(
                        'INSERT INTO report_data (report_id, indicator_id, measured_value, remark) '
                        'VALUES (?, ?, ?, ?)',
                        (id, item['indicator_id'], item.get('measured_value', ''),
                         item.get('remark', ''))
                    )

            # 删除旧的模板字段值
            cursor.execute('DELETE FROM report_field_values WHERE report_id = ?', (id,))

            # 插入新的模板字段值
            for field in template_fields:
                if field.get('field_mapping_id') and field.get('field_value'):
                    cursor.execute(
                        'INSERT INTO report_field_values (report_id, field_mapping_id, field_value) '
                        'VALUES (?, ?, ?)',
                        (id, field['field_mapping_id'], field['field_value'])
                    )

            conn.commit()
            conn.close()

            log_operation('更新报告', f'报告编号:{report["report_number"]}')
            return jsonify({'message': '报告更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': str(e)}), 500

    # GET请求 - 获取报告详情
    report = conn.execute(
        'SELECT r.*, st.name as sample_type_name, st.code as sample_type_code, '
        'c.name as company_name, u.username as creator_name '
        'FROM reports r '
        'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
        'LEFT JOIN companies c ON r.company_id = c.id '
        'LEFT JOIN users u ON r.created_by = u.id '
        'WHERE r.id = ?',
        (id,)
    ).fetchone()

    if not report:
        conn.close()
        return jsonify({'error': '报告不存在'}), 404

    # 获取报告数据
    data = conn.execute(
        'SELECT rd.*, i.name as indicator_name, i.unit, i.limit_value, i.detection_method, '
        'i.group_id, g.name as group_name '
        'FROM report_data rd '
        'LEFT JOIN indicators i ON rd.indicator_id = i.id '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'WHERE rd.report_id = ? '
        'ORDER BY g.sort_order, i.sort_order',
        (id,)
    ).fetchall()

    # 获取模板字段值
    template_fields = []
    if report['template_id']:
        template_fields = conn.execute('''
            SELECT rfv.*, tfm.field_name, tfm.field_display_name
            FROM report_field_values rfv
            LEFT JOIN template_field_mappings tfm ON rfv.field_mapping_id = tfm.id
            WHERE rfv.report_id = ?
        ''', (id,)).fetchall()

    conn.close()

    result = dict(report)
    result['data'] = [dict(row) for row in data]
    result['template_fields'] = [dict(row) for row in template_fields]
    return jsonify(result)

# ==================== 模板导入导出 API ====================
@app.route('/api/templates/export', methods=['POST'])
@admin_required
def api_export_template():
    """导出模板JSON"""
    data = request.json
    sample_type_id = data.get('sample_type_id')

    if not sample_type_id:
        return jsonify({'error': '样品类型ID不能为空'}), 400

    conn = get_db_connection()

    # 获取样品类型信息
    sample_type = conn.execute(
        'SELECT * FROM sample_types WHERE id = ?',
        (sample_type_id,)
    ).fetchone()

    if not sample_type:
        conn.close()
        return jsonify({'error': '样品类型不存在'}), 404

    # 获取关联的检测项目
    template_indicators = conn.execute(
        'SELECT ti.*, i.name as indicator_name, i.unit, i.default_value, i.group_id, '
        'g.name as group_name '
        'FROM template_indicators ti '
        'LEFT JOIN indicators i ON ti.indicator_id = i.id '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'WHERE ti.sample_type_id = ?',
        (sample_type_id,)
    ).fetchall()

    conn.close()

    # 构建导出数据
    export_data = {
        'sample_type': dict(sample_type),
        'indicators': [dict(ti) for ti in template_indicators],
        'export_date': datetime.now().isoformat(),
        'version': '2.0'
    }

    # 保存JSON文件
    os.makedirs('exports', exist_ok=True)
    filename = f"exports/template_{sample_type['code']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.json"

    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(export_data, f, ensure_ascii=False, indent=2)

    log_operation('导出模板', f'导出模板: {sample_type["name"]}')
    return send_file(filename, as_attachment=True, download_name=f"template_{sample_type['code']}.json")

@app.route('/api/templates/import', methods=['POST'])
@admin_required
def api_import_template():
    """导入模板JSON"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    try:
        content = file.read().decode('utf-8')
        data = json.loads(content)

        sample_type_data = data.get('sample_type')
        indicators_data = data.get('indicators', [])

        if not sample_type_data:
            return jsonify({'error': 'JSON格式错误:缺少sample_type'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 检查样品类型是否已存在
        existing = cursor.execute(
            'SELECT id FROM sample_types WHERE code = ?',
            (sample_type_data['code'],)
        ).fetchone()

        if existing:
            conn.close()
            return jsonify({'error': f'样品类型代码 {sample_type_data["code"]} 已存在'}), 400

        # 创建样品类型
        cursor.execute(
            'INSERT INTO sample_types (name, code, description) VALUES (?, ?, ?)',
            (sample_type_data['name'], sample_type_data['code'], sample_type_data.get('description', ''))
        )
        conn.commit()
        sample_type_id = cursor.lastrowid

        # 导入检测项目(需要匹配现有的indicator)
        imported_count = 0
        for item in indicators_data:
            # 查找匹配的indicator
            indicator = cursor.execute(
                'SELECT id FROM indicators WHERE name = ?',
                (item['indicator_name'],)
            ).fetchone()

            if indicator:
                try:
                    cursor.execute(
                        'INSERT INTO template_indicators (sample_type_id, indicator_id, is_required, sort_order) '
                        'VALUES (?, ?, ?, ?)',
                        (sample_type_id, indicator['id'], item.get('is_required', False), item.get('sort_order', 0))
                    )
                    imported_count += 1
                except:
                    pass  # 忽略重复项

        conn.commit()
        conn.close()

        log_operation('导入模板', f'导入模板: {sample_type_data["name"]}, 检测项:{imported_count}')
        return jsonify({
            'message': f'模板导入成功,共导入 {imported_count} 个检测项目',
            'sample_type_id': sample_type_id
        })

    except json.JSONDecodeError:
        return jsonify({'error': 'JSON格式错误'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== 报告导出 API ====================
@app.route('/api/reports/<int:id>/export/excel', methods=['GET'])
@login_required
def api_export_excel(id):
    """导出Excel报告"""
    conn = get_db_connection()

    report = conn.execute(
        'SELECT r.*, st.name as sample_type_name, c.name as company_name '
        'FROM reports r '
        'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
        'LEFT JOIN companies c ON r.company_id = c.id '
        'WHERE r.id = ?',
        (id,)
    ).fetchone()

    if not report:
        conn.close()
        return jsonify({'error': '报告不存在'}), 404

    data = conn.execute(
        'SELECT rd.*, i.name as indicator_name, i.unit, g.name as group_name '
        'FROM report_data rd '
        'LEFT JOIN indicators i ON rd.indicator_id = i.id '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'WHERE rd.report_id = ? '
        'ORDER BY g.sort_order, i.sort_order',
        (id,)
    ).fetchall()

    conn.close()

    # 创建Excel工作簿
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "水质检测报告"

    # 设置样式
    title_font = Font(name='宋体', size=16, bold=True)
    header_font = Font(name='宋体', size=11, bold=True)
    normal_font = Font(name='宋体', size=10)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 标题
    ws.merge_cells('A1:G1')
    title_cell = ws['A1']
    title_cell.value = '水质检测报告'
    title_cell.font = title_font
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30

    # 报告信息
    row = 3
    info_items = [
        ('报告编号', report['report_number']),
        ('样品编号', report['sample_number']),
        ('样品类型', report['sample_type_name']),
        ('委托单位', report['company_name']),
        ('检测日期', report['detection_date']),
        ('检测人员', report['detection_person']),
        ('审核人员', report['review_person'])
    ]

    for label, value in info_items:
        if value:
            ws[f'A{row}'] = label + '：'
            ws[f'B{row}'] = value
            ws[f'A{row}'].font = header_font
            ws.merge_cells(f'B{row}:G{row}')
            row += 1

    row += 1

    # 表头
    headers = ['序号', '检测项目', '单位', '检测结果', '所属分组', '备注']
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col)
        cell.value = header
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    # 数据行
    for idx, item in enumerate(data, start=1):
        row += 1
        row_data = [
            idx,
            item['indicator_name'],
            item['unit'] or '',
            item['measured_value'] or '',
            item['group_name'] or '',
            item['remark'] or ''
        ]

        for col, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row, column=col)
            cell.value = value
            cell.font = normal_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

    # 调整列宽
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 20

    # 保存文件
    os.makedirs('exports', exist_ok=True)
    filename = f"exports/report_{report['report_number']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    wb.save(filename)

    log_operation('导出Excel报告', f'报告编号:{report["report_number"]}')
    return send_file(filename, as_attachment=True, download_name=f"{report['report_number']}.xlsx")

@app.route('/api/reports/<int:id>/export/pdf', methods=['GET'])
@login_required
def api_export_pdf(id):
    """导出PDF报告（先生成Excel再转换为PDF）"""
    import subprocess

    conn = get_db_connection()

    report = conn.execute(
        'SELECT r.*, st.name as sample_type_name, c.name as company_name '
        'FROM reports r '
        'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
        'LEFT JOIN companies c ON r.company_id = c.id '
        'WHERE r.id = ?',
        (id,)
    ).fetchone()

    if not report:
        conn.close()
        return jsonify({'error': '报告不存在'}), 404

    data = conn.execute(
        'SELECT rd.*, i.name as indicator_name, i.unit, g.name as group_name '
        'FROM report_data rd '
        'LEFT JOIN indicators i ON rd.indicator_id = i.id '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'WHERE rd.report_id = ? '
        'ORDER BY g.sort_order, i.sort_order',
        (id,)
    ).fetchall()

    conn.close()

    # 创建Excel工作簿（与Excel导出相同）
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "水质检测报告"

    title_font = Font(name='宋体', size=16, bold=True)
    header_font = Font(name='宋体', size=11, bold=True)
    normal_font = Font(name='宋体', size=10)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    ws.merge_cells('A1:G1')
    title_cell = ws['A1']
    title_cell.value = '水质检测报告'
    title_cell.font = title_font
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30

    row = 3
    info_items = [
        ('报告编号', report['report_number']),
        ('样品编号', report['sample_number']),
        ('样品类型', report['sample_type_name']),
        ('委托单位', report['company_name']),
        ('检测日期', report['detection_date']),
        ('检测人员', report['detection_person']),
        ('审核人员', report['review_person'])
    ]

    for label, value in info_items:
        if value:
            ws[f'A{row}'] = label + '：'
            ws[f'B{row}'] = value
            ws[f'A{row}'].font = header_font
            ws.merge_cells(f'B{row}:G{row}')
            row += 1

    row += 1

    headers = ['序号', '检测项目', '单位', '检测结果', '所属分组', '备注']
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col)
        cell.value = header
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    for idx, item in enumerate(data, start=1):
        row += 1
        row_data = [
            idx,
            item['indicator_name'],
            item['unit'] or '',
            item['measured_value'] or '',
            item['group_name'] or '',
            item['remark'] or ''
        ]
        for col, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row, column=col)
            cell.value = value
            cell.font = normal_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 20

    os.makedirs('exports', exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    xlsx_filename = f"exports/report_{report['report_number']}_{timestamp}.xlsx"
    wb.save(xlsx_filename)

    # 转换为PDF
    try:
        abs_xlsx = os.path.abspath(xlsx_filename)
        abs_outdir = os.path.dirname(abs_xlsx)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', abs_outdir, abs_xlsx],
            capture_output=True, text=True, timeout=30
        )
        pdf_filename = xlsx_filename.replace('.xlsx', '.pdf')
        if result.returncode == 0 and os.path.exists(pdf_filename):
            os.remove(xlsx_filename)
            log_operation('导出PDF报告', f'报告编号:{report["report_number"]}')
            return send_file(pdf_filename, as_attachment=True, download_name=f"{report['report_number']}.pdf")
        else:
            os.remove(xlsx_filename)
            return jsonify({'error': f'PDF转换失败: {result.stderr}'}), 500
    except subprocess.TimeoutExpired:
        if os.path.exists(xlsx_filename):
            os.remove(xlsx_filename)
        return jsonify({'error': 'PDF转换超时'}), 500
    except Exception as e:
        if os.path.exists(xlsx_filename):
            os.remove(xlsx_filename)
        return jsonify({'error': f'PDF转换异常: {str(e)}'}), 500

@app.route('/api/reports/<int:id>/export-simple', methods=['GET'])
@login_required
def api_export_simple_report(id):
    """使用简化模式导出报告"""
    from report_generator import generate_simple_report

    try:
        output_path = generate_simple_report(id)
        log_operation('导出简化报告', f'报告ID: {id}')
        return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/reports/<int:id>/export/word', methods=['GET'])
@login_required
def api_export_word(id):
    """导出Word报告"""
    conn = get_db_connection()

    report = conn.execute(
        'SELECT r.*, st.name as sample_type_name, c.name as company_name '
        'FROM reports r '
        'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
        'LEFT JOIN companies c ON r.company_id = c.id '
        'WHERE r.id = ?',
        (id,)
    ).fetchone()

    if not report:
        conn.close()
        return jsonify({'error': '报告不存在'}), 404

    data = conn.execute(
        'SELECT rd.*, i.name as indicator_name, i.unit, g.name as group_name '
        'FROM report_data rd '
        'LEFT JOIN indicators i ON rd.indicator_id = i.id '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'WHERE rd.report_id = ? '
        'ORDER BY g.sort_order, i.sort_order',
        (id,)
    ).fetchall()

    conn.close()

    # 创建Word文档
    doc = Document()

    # 标题
    title = doc.add_heading('水质检测报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告信息
    info_items = [
        ('报告编号', report['report_number']),
        ('样品编号', report['sample_number']),
        ('样品类型', report['sample_type_name']),
        ('委托单位', report['company_name'] or ''),
        ('检测日期', report['detection_date'] or ''),
        ('检测人员', report['detection_person'] or ''),
        ('审核人员', report['review_person'] or '')
    ]

    for label, value in info_items:
        if value:
            p = doc.add_paragraph()
            p.add_run(f'{label}：').bold = True
            p.add_run(value)

    doc.add_paragraph()

    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # 表头
    headers = ['序号', '检测项目', '单位', '检测结果', '所属分组', '备注']
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for idx, item in enumerate(data, start=1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = item['indicator_name']
        row.cells[2].text = item['unit'] or ''
        row.cells[3].text = item['measured_value'] or ''
        row.cells[4].text = item['group_name'] or ''
        row.cells[5].text = item['remark'] or ''

        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存文件
    os.makedirs('exports', exist_ok=True)
    filename = f"exports/report_{report['report_number']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(filename)

    log_operation('导出Word报告', f'报告编号:{report["report_number"]}')
    return send_file(filename, as_attachment=True, download_name=f"{report['report_number']}.docx")

# ==================== 检测指标导入导出 API ====================
@app.route('/api/indicators/export/excel', methods=['GET'])
@admin_required
def api_export_indicators_excel():
    """导出检测指标到Excel"""
    conn = get_db_connection()

    indicators = conn.execute(
        'SELECT i.*, g.name as group_name '
        'FROM indicators i '
        'LEFT JOIN indicator_groups g ON i.group_id = g.id '
        'ORDER BY i.sort_order, i.name'
    ).fetchall()

    # 关闭连接释放锁
    conn.close()

    # 创建Excel工作簿
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "检测指标"

    # 设置样式
    header_font = Font(name='宋体', size=11, bold=True)
    normal_font = Font(name='宋体', size=10)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 表头
    headers = ['指标名称', '单位', '默认值', '限值', '检测方法', '所属分组', '排序', '备注']
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    # 数据行
    for row_idx, indicator in enumerate(indicators, start=2):
        row_data = [
            indicator['name'],
            indicator['unit'] or '',
            indicator['default_value'] or '',
            indicator['limit_value'] or '',
            indicator['detection_method'] or '',
            indicator['group_name'] or '',
            indicator['sort_order'],
            indicator['remark'] or ''
        ]

        for col, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col)
            cell.value = value
            cell.font = normal_font
            cell.border = border

    # 调整列宽
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 25
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 10
    ws.column_dimensions['H'].width = 30

    # 保存文件
    os.makedirs('exports', exist_ok=True)
    filename = f"exports/indicators_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    wb.save(filename)

    log_operation('导出检测指标', f'导出 {len(indicators)} 个检测指标')

    return send_file(filename, as_attachment=True, download_name='检测指标.xlsx')

@app.route('/api/indicators/import/excel', methods=['POST'])
@admin_required
def api_import_indicators_excel():
    """从Excel导入检测指标"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': '请上传Excel文件(.xlsx 或 .xls)'}), 400

    try:
        # 读取Excel文件
        wb = openpyxl.load_workbook(file)
        ws = wb.active

        conn = get_db_connection()
        cursor = conn.cursor()

        # 获取所有分组，建立名称到ID的映射
        groups = cursor.execute('SELECT id, name FROM indicator_groups').fetchall()
        group_map = {g['name']: g['id'] for g in groups}

        imported_count = 0
        updated_count = 0
        error_rows = []

        # 从第2行开始读取（第1行是表头）
        # 新格式：指标名称、单位、默认值、限值、检测方法、所属分组、排序、备注
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not row[0]:  # 跳过空行
                continue

            name = row[0]
            unit = row[1] or ''
            default_value = row[2] or ''
            limit_value = row[3] or '' if len(row) > 3 else ''
            detection_method = row[4] or '' if len(row) > 4 else ''
            group_name = row[5] or '' if len(row) > 5 else ''
            sort_order = row[6] if len(row) > 6 and row[6] is not None else 0
            remark = row[7] or '' if len(row) > 7 else ''

            # 查找分组ID
            group_id = group_map.get(group_name) if group_name else None

            try:
                # 检查指标是否已存在
                existing = cursor.execute('SELECT id FROM indicators WHERE name = ?', (name,)).fetchone()

                if existing:
                    # 更新现有指标
                    cursor.execute(
                        'UPDATE indicators SET group_id = ?, unit = ?, default_value = ?, limit_value = ?, '
                        'detection_method = ?, remark = ?, sort_order = ? WHERE name = ?',
                        (group_id, unit, default_value, limit_value, detection_method, remark, sort_order, name)
                    )
                    updated_count += 1
                else:
                    # 插入新指标
                    cursor.execute(
                        'INSERT INTO indicators (group_id, name, unit, default_value, limit_value, detection_method, remark, sort_order) '
                        'VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                        (group_id, name, unit, default_value, limit_value, detection_method, remark, sort_order)
                    )
                    imported_count += 1
            except Exception as e:
                error_rows.append(f'第{row_idx}行: {str(e)}')

        conn.commit()
        conn.close()

        log_operation('导入检测指标', f'新增 {imported_count} 个，更新 {updated_count} 个')

        result = {
            'message': f'导入成功！新增 {imported_count} 个指标，更新 {updated_count} 个指标',
            'imported': imported_count,
            'updated': updated_count
        }

        if error_rows:
            result['errors'] = error_rows

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

# ==================== 单个样品类型的检测项目导入导出 API ====================
# 已禁用：使用UI拖拽排序代替Excel导入导出功能
# @app.route('/api/sample-types/<int:sample_type_id>/indicators/export', methods=['GET'])
# @admin_required
# def api_export_sample_type_indicators(sample_type_id):
#     """导出单个样品类型的检测项目到Excel"""
#     conn = get_db_connection()
#     sample_type = conn.execute('SELECT * FROM sample_types WHERE id = ?', (sample_type_id,)).fetchone()
#     if not sample_type:
#         conn.close()
#         return jsonify({'error': '样品类型不存在'}), 404
#     # ... Excel导出代码 ...
#     pass

# 已禁用：使用UI拖拽排序代替Excel导入导出功能
# @app.route('/api/sample-types/<int:sample_type_id>/indicators/import', methods=['POST'])
# @admin_required
# def api_import_sample_type_indicators(sample_type_id):
#     """为单个样品类型导入检测项目 - 严格校验所有字段"""
#     # ... Excel导入代码 ...
#     pass

# 已禁用：使用UI拖拽排序代替Excel导入导出功能
# @app.route('/api/sample-types/import/excel', methods=['POST'])
# @admin_required
# def api_import_sample_types_excel():
#     """从Excel导入样品类型 - 严格校验检测项目信息"""
#     # ... Excel导入代码 ...
#     pass

# ==================== 报告批量导入 API ====================
@app.route('/api/reports/import/excel', methods=['POST'])
@login_required
def api_import_reports_excel():
    """从Excel批量导入报告"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': '请上传Excel文件(.xlsx 或 .xls)'}), 400

    try:
        # 读取Excel文件
        wb = openpyxl.load_workbook(file)
        ws = wb.active

        conn = get_db_connection()
        cursor = conn.cursor()

        # 获取样品类型映射
        sample_types = cursor.execute('SELECT id, name, code FROM sample_types').fetchall()
        sample_type_name_map = {st['name']: st for st in sample_types}
        sample_type_code_map = {st['code']: st for st in sample_types}

        # 获取公司映射
        companies = cursor.execute('SELECT id, name FROM companies').fetchall()
        company_map = {c['name']: c['id'] for c in companies}

        # 获取检测指标映射
        indicators = cursor.execute('SELECT id, name FROM indicators').fetchall()
        indicator_map = {i['name']: i['id'] for i in indicators}

        imported_count = 0
        error_rows = []

        # 读取表头（第1行）
        headers = [cell.value for cell in ws[1]]

        # 查找固定列的索引
        try:
            sample_number_idx = headers.index('样品编号')
            sample_type_idx = headers.index('样品类型')
        except ValueError:
            return jsonify({'error': 'Excel格式错误：必须包含"样品编号"和"样品类型"列'}), 400

        # 可选列
        company_idx = headers.index('委托单位') if '委托单位' in headers else None
        detection_date_idx = headers.index('检测日期') if '检测日期' in headers else None
        detection_person_idx = headers.index('检测人员') if '检测人员' in headers else None
        review_person_idx = headers.index('审核人员') if '审核人员' in headers else None
        remark_idx = headers.index('备注') if '备注' in headers else None

        # 检测指标列（除了固定列之外的列都视为检测指标）
        fixed_cols = {'样品编号', '样品类型', '委托单位', '检测日期', '检测人员', '审核人员', '备注'}
        indicator_cols = [(idx, col) for idx, col in enumerate(headers) if col and col not in fixed_cols]

        # 从第2行开始读取数据
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not row[sample_number_idx]:  # 跳过空行
                continue

            try:
                sample_number = str(row[sample_number_idx]).strip()
                sample_type_value = str(row[sample_type_idx]).strip()

                # 查找样品类型（支持名称或代码）
                sample_type = sample_type_name_map.get(sample_type_value) or sample_type_code_map.get(sample_type_value)

                if not sample_type:
                    error_rows.append(f'第{row_idx}行: 样品类型"{sample_type_value}"不存在')
                    continue

                sample_type_id = sample_type['id']

                # 获取其他字段
                company_id = None
                if company_idx is not None and row[company_idx]:
                    company_name = str(row[company_idx]).strip()
                    company_id = company_map.get(company_name)
                    if not company_id and company_name:
                        # 自动创建公司
                        cursor.execute('INSERT INTO companies (name) VALUES (?)', (company_name,))
                        conn.commit()
                        company_id = cursor.lastrowid
                        company_map[company_name] = company_id

                detection_person = str(row[detection_person_idx]).strip() if detection_person_idx is not None and row[detection_person_idx] else ''
                review_person = str(row[review_person_idx]).strip() if review_person_idx is not None and row[review_person_idx] else ''
                detection_date = str(row[detection_date_idx]) if detection_date_idx is not None and row[detection_date_idx] else None
                remark = str(row[remark_idx]).strip() if remark_idx is not None and row[remark_idx] else ''

                # 生成报告编号
                report_number = f"{sample_number}-{sample_type['code']}"

                # 检查报告编号是否已存在
                existing = cursor.execute('SELECT id FROM reports WHERE report_number = ?', (report_number,)).fetchone()

                if existing:
                    error_rows.append(f'第{row_idx}行: 报告编号"{report_number}"已存在')
                    continue

                # 创建报告
                cursor.execute(
                    'INSERT INTO reports (report_number, sample_number, company_id, sample_type_id, '
                    'detection_person, review_person, detection_date, remark, created_by) '
                    'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                    (report_number, sample_number, company_id, sample_type_id, detection_person,
                     review_person, detection_date, remark, session['user_id'])
                )
                conn.commit()
                report_id = cursor.lastrowid

                # 添加检测数据
                for col_idx, col_name in indicator_cols:
                    if col_name in indicator_map and row[col_idx]:
                        measured_value = str(row[col_idx]).strip()
                        if measured_value:
                            cursor.execute(
                                'INSERT INTO report_data (report_id, indicator_id, measured_value, remark) '
                                'VALUES (?, ?, ?, ?)',
                                (report_id, indicator_map[col_name], measured_value, '')
                            )

                conn.commit()
                imported_count += 1

            except Exception as e:
                error_rows.append(f'第{row_idx}行: {str(e)}')
                continue

        conn.close()

        log_operation('批量导入报告', f'成功导入 {imported_count} 份报告')

        result = {
            'message': f'导入成功！共导入 {imported_count} 份报告',
            'imported': imported_count
        }

        if error_rows:
            result['errors'] = error_rows

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

@app.route('/api/reports/export/template', methods=['GET'])
@login_required
def api_export_reports_template():
    """导出报告导入模板Excel"""
    sample_type_id = request.args.get('sample_type_id')

    if not sample_type_id:
        return jsonify({'error': '请指定样品类型'}), 400

    conn = get_db_connection()

    # 获取样品类型信息
    sample_type = conn.execute('SELECT * FROM sample_types WHERE id = ?', (sample_type_id,)).fetchone()

    if not sample_type:
        conn.close()
        return jsonify({'error': '样品类型不存在'}), 404

    # 获取该样品类型的检测指标
    indicators = conn.execute(
        'SELECT i.name, i.unit '
        'FROM template_indicators ti '
        'LEFT JOIN indicators i ON ti.indicator_id = i.id '
        'WHERE ti.sample_type_id = ? '
        'ORDER BY ti.sort_order',
        (sample_type_id,)
    ).fetchall()

    conn.close()

    # 创建Excel工作簿
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "检测数据"

    # 设置样式
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    subheader_fill = PatternFill(start_color="B4C7E7", end_color="B4C7E7", fill_type="solid")
    subheader_font = Font(name='宋体', size=11, bold=True)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 简化格式表头：检测项目、单位、样品数据列
    headers = ['检测项目', '单位']
    sample_numbers = ['样品编号1*', '样品编号2', '样品编号3']

    for col, header in enumerate(headers + sample_numbers, start=1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        if col <= len(headers):
            cell.fill = subheader_fill
            cell.font = subheader_font
        else:
            cell.fill = header_fill
            cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    # 添加检测指标行
    for row_idx, indicator in enumerate(indicators, start=2):
        # 检测项目（A列）
        cell = ws.cell(row=row_idx, column=1)
        cell.value = indicator['name']
        cell.alignment = Alignment(horizontal='left', vertical='center')
        cell.border = border

        # 单位（B列）
        cell = ws.cell(row=row_idx, column=2)
        cell.value = indicator['unit'] or ''
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

        # 样品数据列（C列起）留空
        for col_idx in range(3, 3 + len(sample_numbers)):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')

    # 调整列宽
    ws.column_dimensions['A'].width = 20  # 检测项目
    ws.column_dimensions['B'].width = 12  # 单位
    for col_letter in ['C', 'D', 'E', 'F', 'G']:
        ws.column_dimensions[col_letter].width = 15

    # 冻结首行和前2列
    ws.freeze_panes = 'C2'

    # 保存文件
    os.makedirs('exports', exist_ok=True)
    filename = f"exports/report_template_{sample_type['code']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    wb.save(filename)

    log_operation('导出报告模板', f'导出样品类型:{sample_type["name"]}')
    return send_file(filename, as_attachment=True, download_name=f'报告导入模板_{sample_type["code"]}.xlsx')

# ==================== 数据备份与恢复 API ====================
@app.route('/api/backup/create', methods=['POST'])
@admin_required
def api_create_backup():
    """创建数据备份"""
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_dir = f'backups/backup_{timestamp}'
        os.makedirs(backup_dir, exist_ok=True)

        # 备份数据库文件
        if os.path.exists(DATABASE_PATH):
            shutil.copy2(DATABASE_PATH, f'{backup_dir}/water_quality_v2.db')

        # 创建备份信息文件
        backup_info = {
            'backup_time': datetime.now().isoformat(),
            'backup_by': session.get('username', 'unknown'),
            'version': '2.0'
        }

        with open(f'{backup_dir}/backup_info.json', 'w', encoding='utf-8') as f:
            json.dump(backup_info, f, ensure_ascii=False, indent=2)

        log_operation('创建数据备份', f'备份目录:{backup_dir}')
        return jsonify({'message': '备份创建成功', 'backup_dir': backup_dir})

    except Exception as e:
        return jsonify({'error': f'备份失败: {str(e)}'}), 500

@app.route('/api/backup/list', methods=['GET'])
@admin_required
def api_list_backups():
    """获取备份列表"""
    try:
        backups = []
        backup_base = 'backups'

        if os.path.exists(backup_base):
            for backup_name in os.listdir(backup_base):
                backup_path = os.path.join(backup_base, backup_name)
                if os.path.isdir(backup_path):
                    info_file = os.path.join(backup_path, 'backup_info.json')
                    if os.path.exists(info_file):
                        with open(info_file, 'r', encoding='utf-8') as f:
                            info = json.load(f)
                            info['name'] = backup_name
                            info['path'] = backup_path
                            backups.append(info)

        backups.sort(key=lambda x: x['backup_time'], reverse=True)
        return jsonify(backups)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/backup/restore', methods=['POST'])
@admin_required
def api_restore_backup():
    """恢复数据备份"""
    data = request.json
    backup_name = data.get('backup_name')

    if not backup_name:
        return jsonify({'error': '备份名称不能为空'}), 400

    backup_path = os.path.join('backups', backup_name)

    if not os.path.exists(backup_path):
        return jsonify({'error': '备份不存在'}), 404

    try:
        # 备份当前数据库(防止恢复失败)
        if os.path.exists(DATABASE_PATH):
            shutil.copy2(DATABASE_PATH, f'{DATABASE_PATH}.before_restore')

        # 恢复数据库文件
        backup_db = os.path.join(backup_path, 'water_quality_v2.db')
        if os.path.exists(backup_db):
            shutil.copy2(backup_db, DATABASE_PATH)

        log_operation('恢复数据备份', f'恢复备份:{backup_name}')
        return jsonify({'message': '数据恢复成功'})

    except Exception as e:
        # 恢复失败,回滚
        if os.path.exists(f'{DATABASE_PATH}.before_restore'):
            shutil.copy2(f'{DATABASE_PATH}.before_restore', DATABASE_PATH)
        return jsonify({'error': f'恢复失败: {str(e)}'}), 500

# ==================== 操作日志 API ====================
@app.route('/api/logs', methods=['GET'])
@login_required
def api_logs():
    """获取操作日志"""
    limit = int(request.args.get('limit', 100))
    offset = int(request.args.get('offset', 0))
    user_id = request.args.get('user_id')
    operation_type = request.args.get('operation_type')

    logs = get_operation_logs(limit, offset, user_id, operation_type)

    return jsonify(logs)

# ==================== 报告模版管理 API ====================
@app.route('/api/report-templates', methods=['GET'])
@login_required
def api_report_templates():
    """获取报告模版列表"""
    conn = get_db_connection()

    templates = conn.execute(
        'SELECT t.*, st.name as sample_type_name '
        'FROM excel_report_templates t '
        'LEFT JOIN sample_types st ON t.sample_type_id = st.id '
        'WHERE t.is_active = 1 '
        'ORDER BY t.created_at DESC'
    ).fetchall()

    conn.close()

    return jsonify([dict(t) for t in templates])

@app.route('/api/report-templates/import', methods=['POST'])
@admin_required
def api_import_report_template():
    """导入报告模版"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': '只支持Excel文件'}), 400

    # 获取表单数据
    template_name = request.form.get('name')
    sample_type_id = request.form.get('sample_type_id')
    description = request.form.get('description', '')

    if not template_name:
        return jsonify({'error': '模版名称不能为空'}), 400

    try:
        # 保存上传的文件
        os.makedirs('templates/excel_reports', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{template_name}_{timestamp}.xlsx"
        file_path = os.path.join('templates/excel_reports', filename)
        file.save(file_path)

        # 读取Excel文件
        wb = openpyxl.load_workbook(file_path)

        # 保存到数据库
        conn = get_db_connection()
        cursor = conn.cursor()

        # 检查模板名称是否已存在，如果存在则自动添加序号
        final_template_name = template_name
        existing = conn.execute(
            'SELECT COUNT(*) as count FROM excel_report_templates WHERE name = ?',
            (template_name,)
        ).fetchone()

        if existing['count'] > 0:
            # 查找所有相似名称的模板
            similar = conn.execute(
                'SELECT name FROM excel_report_templates WHERE name LIKE ?',
                (f'{template_name}%',)
            ).fetchall()

            # 找出最大的序号
            max_num = 0
            for row in similar:
                name = row['name']
                # 尝试提取末尾的数字
                import re
                match = re.search(r'_(\d+)$', name)
                if match:
                    num = int(match.group(1))
                    if num > max_num:
                        max_num = num

            # 使用下一个序号
            final_template_name = f"{template_name}_{max_num + 1}"

        cursor.execute(
            'INSERT INTO excel_report_templates (name, sample_type_id, description, template_file_path) '
            'VALUES (?, ?, ?, ?)',
            (final_template_name, sample_type_id if sample_type_id else None, description, file_path)
        )

        template_id = cursor.lastrowid

        # 分析工作表结构
        for index, sheet_name in enumerate(wb.sheetnames):
            sheet_type = identify_sheet_type(sheet_name)
            page_number = extract_page_number(sheet_name)

            cursor.execute(
                'INSERT INTO template_sheet_configs '
                '(template_id, sheet_name, sheet_index, sheet_type, page_number) '
                'VALUES (?, ?, ?, ?, ?)',
                (template_id, sheet_name, index, sheet_type, page_number)
            )

        # 解析模板字段（带有[]、()、;标记的单元格）
        from template_field_parser import TemplateFieldParser

        field_count = 0
        try:
            fields = TemplateFieldParser.extract_template_fields(file_path)

            for field in fields:
                # 插入字段映射
                cursor.execute(
                    '''INSERT INTO template_field_mappings
                       (template_id, field_name, field_display_name, field_type,
                        sheet_name, cell_address, placeholder, default_value, is_required,
                        original_cell_text, field_code, is_reference, column_mapping)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                    (template_id,
                     field['field_name'],
                     field['display_name'],
                     field.get('field_type', 'text'),  # 使用解析得到的字段类型
                     field['sheet_name'],
                     field['cell_address'],
                     field.get('placeholder', ''),
                     field.get('default_value', ''),
                     1 if field.get('is_required', True) else 0,
                     field.get('original_value', ''),  # 保存原始单元格文本
                     field.get('field_code'),  # 保存字段代号（如 #report_no）
                     1 if field.get('is_reference', False) else 0,  # 是否为引用字段
                     field.get('column_mapping', ''))  # 检测数据列映射
                )
                field_count += 1
        except Exception as e:
            print(f"字段解析警告: {e}")
            # 字段解析失败不影响模板导入

        conn.commit()
        conn.close()

        log_operation('导入报告模版', f'导入模版: {template_name}, 解析字段: {field_count}个')

        return jsonify({
            'id': template_id,
            'message': '模版导入成功',
            'sheet_count': len(wb.sheetnames),
            'field_count': field_count
        }), 201

    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

@app.route('/api/report-templates/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_report_template_detail(id):
    """获取、修改或删除报告模版"""
    conn = get_db_connection()

    if request.method == 'PUT':
        # 仅管理员可修改
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        data = request.json
        name = data.get('name')
        sample_type_id = data.get('sample_type_id')
        description = data.get('description', '')

        if not name:
            return jsonify({'error': '模版名称不能为空'}), 400

        # 检查是否存在同名的其他模版（包括已删除的模版，因为UNIQUE约束对所有行生效）
        existing = conn.execute(
            'SELECT id, is_active, name FROM excel_report_templates WHERE name = ? AND id != ?',
            (name, id)
        ).fetchone()

        if existing:
            if existing['is_active']:
                conn.close()
                return jsonify({'error': '模版名称已存在，请使用其他名称'}), 400
            else:
                # 如果是已删除的模版，自动重命名它以释放该名称
                import datetime
                timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
                new_name_for_deleted = f"{existing['name']}_已删除_{timestamp}"
                conn.execute(
                    'UPDATE excel_report_templates SET name = ? WHERE id = ?',
                    (new_name_for_deleted, existing['id'])
                )
                conn.commit()

        try:
            conn.execute(
                'UPDATE excel_report_templates SET name = ?, sample_type_id = ?, description = ? WHERE id = ?',
                (name, sample_type_id, description, id)
            )
            conn.commit()

            log_operation('修改报告模版', f'修改模版: {name}', conn=conn)
            conn.close()

            return jsonify({'message': '模版更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': f'更新失败: {str(e)}'}), 500

    if request.method == 'DELETE':
        # 仅管理员可删除
        if session.get('role') != 'admin':
            return jsonify({'error': '需要管理员权限'}), 403

        # 获取模版信息
        template = conn.execute(
            'SELECT * FROM excel_report_templates WHERE id = ?',
            (id,)
        ).fetchone()

        if not template:
            conn.close()
            return jsonify({'error': '模版不存在'}), 404

        # 删除文件
        if template['template_file_path'] and os.path.exists(template['template_file_path']):
            try:
                os.remove(template['template_file_path'])
            except Exception as e:
                print(f"删除模版文件失败: {e}")

        # 软删除（设置is_active=0）
        conn.execute('UPDATE excel_report_templates SET is_active = 0 WHERE id = ?', (id,))
        conn.commit()
        conn.close()

        log_operation('删除报告模版', f'删除模版: {template["name"]}')

        return jsonify({'message': '模版删除成功'})

    # GET请求 - 获取模版详情
    template = conn.execute(
        'SELECT * FROM excel_report_templates WHERE id = ?',
        (id,)
    ).fetchone()

    if not template:
        conn.close()
        return jsonify({'error': '模版不存在'}), 404

    # 获取工作表配置
    sheets = conn.execute(
        'SELECT * FROM template_sheet_configs WHERE template_id = ? ORDER BY sheet_index',
        (id,)
    ).fetchall()

    # 获取字段映射
    fields = conn.execute(
        'SELECT * FROM template_field_mappings WHERE template_id = ?',
        (id,)
    ).fetchall()

    conn.close()

    return jsonify({
        'template': dict(template),
        'sheets': [dict(s) for s in sheets],
        'fields': [dict(f) for f in fields]
    })

@app.route('/api/report-templates/<int:id>/fields', methods=['GET', 'POST'])
@admin_required
def api_template_fields(id):
    """获取或添加模版字段映射"""
    conn = get_db_connection()

    if request.method == 'POST':
        data = request.json
        field_name = data.get('field_name')
        field_type = data.get('field_type')
        sheet_name = data.get('sheet_name')
        cell_address = data.get('cell_address')
        start_row = data.get('start_row')
        start_col = data.get('start_col')
        description = data.get('description', '')
        is_required = data.get('is_required', False)
        default_value = data.get('default_value', '')

        if not all([field_name, field_type, sheet_name]):
            conn.close()
            return jsonify({'error': '缺少必填字段'}), 400

        cursor = conn.cursor()
        cursor.execute(
            'INSERT INTO template_field_mappings '
            '(template_id, field_name, field_type, sheet_name, cell_address, '
            'start_row, start_col, description, is_required, default_value) '
            'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
            (id, field_name, field_type, sheet_name, cell_address,
             start_row, start_col, description, is_required, default_value)
        )

        field_id = cursor.lastrowid
        conn.commit()
        conn.close()

        log_operation('添加模版字段映射', f'模版ID: {id}, 字段: {field_name}')

        return jsonify({'id': field_id, 'message': '字段映射添加成功'}), 201

    # GET请求
    fields = conn.execute(
        'SELECT * FROM template_field_mappings WHERE template_id = ? ORDER BY id',
        (id,)
    ).fetchall()

    conn.close()

    return jsonify([dict(f) for f in fields])

@app.route('/api/template-fields/<int:field_id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_template_field_detail(field_id):
    """获取、更新或删除单个模板字段"""
    conn = get_db_connection()

    if request.method == 'GET':
        # 获取字段详情
        field = conn.execute(
            'SELECT * FROM template_field_mappings WHERE id = ?',
            (field_id,)
        ).fetchone()

        conn.close()

        if not field:
            return jsonify({'error': '字段不存在'}), 404

        return jsonify(dict(field))

    if request.method == 'PUT':
        # 仅管理员可修改
        if session.get('role') != 'admin':
            conn.close()
            return jsonify({'error': '需要管理员权限'}), 403

        # 检查字段是否存在
        field = conn.execute(
            'SELECT * FROM template_field_mappings WHERE id = ?',
            (field_id,)
        ).fetchone()

        if not field:
            conn.close()
            return jsonify({'error': '字段不存在'}), 404

        data = request.json
        field_name = data.get('field_name')
        field_display_name = data.get('field_display_name', '')
        field_type = data.get('field_type')
        sheet_name = data.get('sheet_name')
        cell_address = data.get('cell_address', '')
        placeholder = data.get('placeholder', '')
        default_value = data.get('default_value', '')
        is_required = data.get('is_required', 0)
        description = data.get('description', '')

        if not all([field_name, field_type, sheet_name]):
            conn.close()
            return jsonify({'error': '缺少必填字段'}), 400

        try:
            conn.execute(
                '''UPDATE template_field_mappings
                   SET field_name = ?, field_display_name = ?, field_type = ?,
                       sheet_name = ?, cell_address = ?, placeholder = ?,
                       default_value = ?, is_required = ?, description = ?
                   WHERE id = ?''',
                (field_name, field_display_name, field_type, sheet_name,
                 cell_address, placeholder, default_value, is_required,
                 description, field_id)
            )
            conn.commit()

            log_operation('更新模板字段', f'字段ID: {field_id}, 字段名: {field_name}', conn=conn)
            conn.close()

            return jsonify({'message': '字段更新成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': f'更新失败: {str(e)}'}), 500

    if request.method == 'DELETE':
        # 仅管理员可删除
        if session.get('role') != 'admin':
            conn.close()
            return jsonify({'error': '需要管理员权限'}), 403

        # 检查字段是否存在
        field = conn.execute(
            'SELECT * FROM template_field_mappings WHERE id = ?',
            (field_id,)
        ).fetchone()

        if not field:
            conn.close()
            return jsonify({'error': '字段不存在'}), 404

        try:
            conn.execute('DELETE FROM template_field_mappings WHERE id = ?', (field_id,))
            conn.commit()

            log_operation('删除模板字段', f'字段ID: {field_id}, 字段名: {field["field_name"]}', conn=conn)
            conn.close()

            return jsonify({'message': '字段删除成功'})
        except Exception as e:
            conn.close()
            return jsonify({'error': f'删除失败: {str(e)}'}), 500

@app.route('/api/report-templates/<int:id>/export-config', methods=['GET'])
@login_required
def api_export_template_config(id):
    """导出模板配置为Excel文件"""
    from template_config_excel import TemplateConfigExcel

    try:
        output_path = TemplateConfigExcel.export_template_config(id)

        log_operation('导出模板配置', f'导出模板ID: {id}的配置')

        return send_file(
            output_path,
            as_attachment=True,
            download_name=os.path.basename(output_path),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except ValueError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/report-templates/<int:id>/import-config', methods=['POST'])
@admin_required
def api_import_template_config(id):
    """从Excel文件导入模板配置"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': '只支持Excel文件'}), 400

    from template_config_excel import TemplateConfigExcel

    try:
        # 保存上传的文件到临时位置
        os.makedirs('temp/config_imports', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_path = f'temp/config_imports/config_import_{timestamp}.xlsx'
        file.save(temp_path)

        # 导入配置
        result = TemplateConfigExcel.import_template_config(id, temp_path)

        # 删除临时文件
        try:
            os.remove(temp_path)
        except:
            pass

        log_operation('导入模板配置', f'导入模板ID: {id}的配置，共{result["inserted_count"]}个字段')

        return jsonify(result), 200

    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

def identify_sheet_type(sheet_name):
    """识别工作表类型"""
    sheet_name_lower = sheet_name.lower()

    if '1' in sheet_name or 'cover' in sheet_name_lower or '封面' in sheet_name:
        return 'cover'
    elif '2' in sheet_name or 'info' in sheet_name_lower or '信息' in sheet_name:
        return 'info'
    elif any(x in sheet_name for x in ['3', '4']) or 'data' in sheet_name_lower or '数据' in sheet_name:
        return 'data'
    elif '5' in sheet_name or 'note' in sheet_name_lower or '说明' in sheet_name:
        return 'conclusion'
    else:
        return 'other'

def extract_page_number(sheet_name):
    """从工作表名称提取页码"""
    import re
    match = re.search(r'\d+', sheet_name)
    return int(match.group()) if match else 0

# ==================== 页面路由 ====================
@app.route('/login')
def login_page():
    """登录页面"""
    return render_template('login.html')

@app.route('/')
def index():
    """主页面 - 需要登录"""
    # 检查是否已登录
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('index_v2.html')

@app.route('/sample-types-manager')
def sample_types_manager():
    """样品类型管理专项页面"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('sample_types_manager.html')

@app.route('/indicators-manager')
def indicators_manager():
    """检测指标管理专项页面"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('indicators_manager.html')

@app.route('/report-template-manager')
def report_template_manager():
    """报告模版管理专项页面"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('report_template_manager.html')

@app.route('/raw-data-manager')
def raw_data_manager():
    """原始数据管理专项页面"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('raw_data_manager.html')

@app.route('/customers-manager')
def customers_manager():
    """客户管理专项页面"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('customers_manager.html')

@app.route('/report-templates')
def report_templates_page():
    """报告模版管理页面（新版，使用别名）"""
    if 'user_id' not in session:
        return render_template('login.html')
    return render_template('report_template_manager.html')

# ==================== 已禁用：模板配置编辑器页面路由 ====================
# 模板配置编辑器已被禁用，因为模板配置功能已整合到样品类型管理中
# 详见 REMOVED_FEATURES.md

# @app.route('/template-config-editor')
# def template_config_editor():
#     """模板配置编辑器页面"""
#     if 'user_id' not in session:
#         return render_template('login.html')
#     return render_template('template_config_editor.html')

# ==================== 新增API接口 ====================

@app.route('/api/field-code-reference', methods=['GET'])
@login_required
def api_download_field_code_reference():
    """下载字段代号使用说明文档"""
    from field_code_mapping import FieldCodeMapping
    import io

    try:
        # 生成文档内容
        doc_content = FieldCodeMapping.generate_documentation()

        # 创建文本文件
        output = io.BytesIO()
        output.write(doc_content.encode('utf-8'))
        output.seek(0)

        log_operation('下载字段代号文档', '字段代号使用说明')

        return send_file(
            output,
            mimetype='text/plain; charset=utf-8',
            as_attachment=True,
            download_name='Excel模板字段代号使用说明.txt'
        )
    except Exception as e:
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/api/download-example-template', methods=['GET'])
@login_required
def api_download_example_template():
    """下载Excel模板示例文件"""
    try:
        template_path = 'template_examples/水质检测报告模板示例.xlsx'

        # 如果文件不存在，先生成
        if not os.path.exists(template_path):
            from generate_example_template import create_example_template
            create_example_template()

        log_operation('下载模板示例', 'Excel模板示例文件')

        return send_file(
            template_path,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='水质检测报告模板示例.xlsx'
        )
    except Exception as e:
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/api/export-report-template/<int:template_id>', methods=['GET'])
@login_required
def api_export_report_template(template_id):
    """导出报告填写模板"""
    from report_template_exporter import export_report_template

    try:
        output_path = export_report_template(template_id)
        log_operation('导出报告填写模板', f'模板ID: {template_id}')
        return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/export-sample-type-template/<int:sample_type_id>', methods=['GET'])
@login_required
def api_export_sample_type_template(sample_type_id):
    """导出样品类型检测模板"""
    from sample_type_exporter import export_sample_type_template

    try:
        output_path = export_sample_type_template(sample_type_id)
        log_operation('导出检测项目模板', f'样品类型ID: {sample_type_id}')
        return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/import-report-info', methods=['POST'])
@login_required
def api_import_report_info():
    """导入报告基本信息"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    try:
        # 保存上传的文件
        os.makedirs('temp/imports', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_path = f'temp/imports/report_info_{timestamp}.xlsx'
        file.save(temp_path)

        # 读取Excel文件
        wb = openpyxl.load_workbook(temp_path, data_only=True)
        ws = wb['报告基本信息']

        conn = get_db_connection()
        cursor = conn.cursor()
        created_count = 0

        # 获取字段名称（第一列，从第2行开始）
        field_names = []
        row_idx = 2
        while True:
            field_name = ws.cell(row_idx, 1).value
            if field_name is None:
                break
            field_names.append(field_name.replace('*', '').strip())
            row_idx += 1

        # 处理每一列数据（从第2列开始）
        col_idx = 2
        while True:
            sample_number = ws.cell(1, col_idx).value
            if sample_number is None or str(sample_number).strip() == '':
                break

            sample_number = str(sample_number).replace('*', '').strip()

            # 读取该列的所有数据
            report_data = {}
            for i, field_name in enumerate(field_names, start=2):
                cell_value = ws.cell(i, col_idx).value
                report_data[field_name] = cell_value if cell_value is not None else ''

            # 创建报告记录（简化版本，实际需要根据模板字段创建）
            # 这里暂时创建基本报告记录
            report_number = f"RPT{datetime.now().strftime('%Y%m%d%H%M%S')}{created_count+1:03d}"

            # 注意：这里需要根据实际的模板字段映射来创建报告
            # 暂时先创建一个简单的占位实现
            created_count += 1
            col_idx += 1

        conn.commit()
        conn.close()

        # 删除临时文件
        try:
            os.remove(temp_path)
        except:
            pass

        log_operation('导入报告基本信息', f'成功导入 {created_count} 份报告')

        return jsonify({
            'message': '导入成功',
            'created_count': created_count
        }), 200

    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

@app.route('/api/import-detection-data', methods=['POST'])
@login_required
def api_import_detection_data():
    """导入检测项目数据"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    try:
        # 保存上传的文件
        os.makedirs('temp/imports', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_path = f'temp/imports/detection_data_{timestamp}.xlsx'
        file.save(temp_path)

        # 读取Excel文件
        wb = openpyxl.load_workbook(temp_path, data_only=True)
        ws = wb['检测数据']

        conn = get_db_connection()
        cursor = conn.cursor()
        updated_count = 0

        # 获取检测项目列表（第2列，从第2行开始）
        indicators = []
        row_idx = 2
        while True:
            indicator_name = ws.cell(row_idx, 2).value
            if indicator_name is None:
                break
            indicators.append(indicator_name)
            row_idx += 1

        # 处理每一列样品数据（从第6列开始）
        col_idx = 6
        while True:
            sample_number = ws.cell(1, col_idx).value
            if sample_number is None or str(sample_number).strip() == '':
                break

            sample_number = str(sample_number).replace('*', '').strip()

            # 查找对应的报告
            report = conn.execute(
                'SELECT id FROM reports WHERE sample_number = ?',
                (sample_number,)
            ).fetchone()

            if report:
                report_id = report['id']

                # 读取该列的检测数据
                for i, indicator_name in enumerate(indicators, start=2):
                    measured_value = ws.cell(i, col_idx).value
                    if measured_value is not None and str(measured_value).strip() != '':
                        # 查找指标ID
                        indicator = conn.execute(
                            'SELECT id FROM indicators WHERE name = ?',
                            (indicator_name,)
                        ).fetchone()

                        if indicator:
                            # 更新或插入检测数据
                            cursor.execute('''
                                INSERT OR REPLACE INTO report_data (report_id, indicator_id, measured_value)
                                VALUES (?, ?, ?)
                            ''', (report_id, indicator['id'], str(measured_value)))

                updated_count += 1

            col_idx += 1

        conn.commit()
        conn.close()

        # 删除临时文件
        try:
            os.remove(temp_path)
        except:
            pass

        log_operation('导入检测数据', f'成功更新 {updated_count} 份报告的检测数据')

        return jsonify({
            'message': '导入成功',
            'updated_count': updated_count
        }), 200

    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500

@app.route('/api/validate-report-excel', methods=['POST'])
@login_required
def api_validate_report_excel():
    """验证上传的Excel文件格式"""
    try:
        # 获取参数
        template_id = request.form.get('template_id')
        sample_type_id = request.form.get('sample_type_id')

        if not template_id or not sample_type_id:
            return jsonify({'valid': False, 'errors': ['缺少必要参数：template_id 或 sample_type_id']}), 400

        # 获取上传的文件
        if 'template_excel' not in request.files or 'detection_excel' not in request.files:
            return jsonify({'valid': False, 'errors': ['缺少Excel文件']}), 400

        template_file = request.files['template_excel']
        detection_file = request.files['detection_excel']

        if not template_file.filename or not detection_file.filename:
            return jsonify({'valid': False, 'errors': ['文件名为空']}), 400

        validation_errors = []
        validation_warnings = []

        # 保存临时文件
        os.makedirs('temp/validate', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        template_path = f'temp/validate/template_{timestamp}.xlsx'
        detection_path = f'temp/validate/detection_{timestamp}.xlsx'

        template_file.save(template_path)
        detection_file.save(detection_path)

        # 验证报告模板Excel
        try:
            template_wb = openpyxl.load_workbook(template_path, data_only=True)

            # 检查是否有"报告基本信息"工作表
            if '报告基本信息' not in template_wb.sheetnames:
                validation_warnings.append('报告模板Excel中没有找到"报告基本信息"工作表，将无法自动填充基本信息')
            else:
                info_ws = template_wb['报告基本信息']
                # 验证基本信息格式
                expected_fields = ['样品编号', '委托单位', '检测日期', '检测人员', '审核人员', '备注']
                found_fields = []
                for row_idx in range(2, min(20, info_ws.max_row + 1)):
                    field_name = info_ws.cell(row_idx, 1).value
                    if field_name:
                        found_fields.append(str(field_name).strip())

                missing_fields = [f for f in expected_fields if f not in found_fields]
                if missing_fields:
                    validation_warnings.append(f'报告基本信息中缺少以下字段：{", ".join(missing_fields)}')

            template_wb.close()

        except Exception as e:
            validation_errors.append(f'报告模板Excel格式错误：{str(e)}')

        # 验证检测数据Excel
        try:
            detection_wb = openpyxl.load_workbook(detection_path, data_only=True)
            detection_ws = detection_wb.active

            # 检查是否有数据
            if detection_ws.max_row < 2:
                validation_errors.append('检测数据Excel中没有数据行')

            # 检查A列（第1列）是否有指标名称（简化格式）
            indicator_count = 0
            for row_idx in range(2, min(100, detection_ws.max_row + 1)):
                indicator_name = detection_ws.cell(row_idx, 1).value
                if indicator_name and str(indicator_name).strip():
                    indicator_count += 1

            if indicator_count == 0:
                validation_errors.append('检测数据Excel的A列（检测项目列）没有找到任何指标')

            # 检查C列（第3列）是否有样品编号
            sample_number = detection_ws.cell(1, 3).value
            if not sample_number or str(sample_number).strip() == '':
                validation_warnings.append('检测数据Excel的C列（样品数据列）标题为空，请在首行C列填写样品编号')

            detection_wb.close()

        except Exception as e:
            validation_errors.append(f'检测数据Excel格式错误：{str(e)}')

        # 清理临时文件
        try:
            os.remove(template_path)
            os.remove(detection_path)
        except:
            pass

        # 返回验证结果
        is_valid = len(validation_errors) == 0

        return jsonify({
            'valid': is_valid,
            'errors': validation_errors,
            'warnings': validation_warnings
        }), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'valid': False,
            'errors': [f'验证失败：{str(e)}']
        }), 500

@app.route('/api/parse-report-excel', methods=['POST'])
@login_required
def api_parse_report_excel():
    """解析上传的报告模板Excel和检测数据Excel"""
    try:
        # 获取参数
        template_id = request.form.get('template_id')
        sample_type_id = request.form.get('sample_type_id')

        if not template_id or not sample_type_id:
            return jsonify({'error': '缺少必要参数'}), 400

        # 获取上传的文件
        if 'template_excel' not in request.files or 'detection_excel' not in request.files:
            return jsonify({'error': '缺少Excel文件'}), 400

        template_file = request.files['template_excel']
        detection_file = request.files['detection_excel']

        if not template_file.filename or not detection_file.filename:
            return jsonify({'error': '文件名为空'}), 400

        # 保存临时文件
        os.makedirs('temp/parse', exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        template_path = f'temp/parse/template_{timestamp}.xlsx'
        detection_path = f'temp/parse/detection_{timestamp}.xlsx'

        template_file.save(template_path)
        detection_file.save(detection_path)

        # 解析报告模板Excel
        template_wb = openpyxl.load_workbook(template_path, data_only=True)

        # 解析检测数据Excel
        detection_wb = openpyxl.load_workbook(detection_path, data_only=True)
        detection_ws = detection_wb.active

        conn = get_db_connection()
        cursor = conn.cursor()

        # 解析基本信息（从报告模板Excel中提取）
        basic_info = {}

        # 定义基本信息字段，避免与模板字段混淆
        basic_info_fields = {
            '样品编号': 'sample_number',
            '委托单位': 'company_name',
            '检测日期': 'detection_date',
            '检测人员': 'detection_person',
            '审核人员': 'review_person',
            '备注': 'remark'
        }

        # 尝试从第一个工作表读取基本信息
        if '报告基本信息' in template_wb.sheetnames:
            info_ws = template_wb['报告基本信息']
            # 读取字段名和值（假设格式：列A是字段名，列B是值）
            row_idx = 2
            while row_idx <= info_ws.max_row:
                field_name = info_ws.cell(row_idx, 1).value
                field_value = info_ws.cell(row_idx, 2).value

                if not field_name:
                    break

                field_name_str = str(field_name).strip()

                # 只处理基本信息字段
                if field_name_str in basic_info_fields:
                    standard_field = basic_info_fields[field_name_str]
                    basic_info[standard_field] = str(field_value) if field_value else ''

                row_idx += 1

        # 解析模板字段
        template_fields = []
        fields_result = cursor.execute('''
            SELECT fm.id, fm.field_name, fm.field_name as field_display_name, fm.field_type,
                   fm.is_required, fm.sheet_name, fm.cell_address
            FROM template_field_mappings fm
            WHERE fm.template_id = ?
            ORDER BY fm.id
        ''', (template_id,)).fetchall()

        for field in fields_result:
            field_data = {
                'field_mapping_id': field[0],
                'field_name': field[2] or field[1],
                'field_type': field[3],
                'is_required': field[4],
                'field_value': ''
            }

            # 尝试从Excel中读取字段值
            if field[5] and field[6]:  # sheet_name 和 cell_address
                try:
                    if field[5] in template_wb.sheetnames:
                        ws = template_wb[field[5]]
                        cell_value = ws[field[6]].value
                        if cell_value:
                            field_data['field_value'] = str(cell_value)
                except:
                    pass

            template_fields.append(field_data)

        # 解析检测数据
        detection_data = []

        # 获取该样品类型的指标列表
        indicators_result = cursor.execute('''
            SELECT i.id, i.name, i.unit, i.default_value, i.limit_value,
                   ig.name as group_name, ti.sort_order
            FROM template_indicators ti
            JOIN indicators i ON ti.indicator_id = i.id
            LEFT JOIN indicator_groups ig ON i.group_id = ig.id
            WHERE ti.sample_type_id = ?
            ORDER BY ti.sort_order
        ''', (sample_type_id,)).fetchall()

        # 从检测数据Excel中读取数据
        # 简化格式：第1列检测项目，第2列单位（参考），第3列及以后为样品检测数据
        indicator_name_col = 1  # 指标名称列（A列）
        unit_col = 2            # 单位列（B列，参考用）
        value_col = 3           # 检测值列（C列，第一个样品的数据）

        for indicator in indicators_result:
            indicator_id = indicator[0]
            indicator_name = indicator[1]
            unit = indicator[2]
            default_value = indicator[3]
            limit_value = indicator[4]
            group_name = indicator[5]

            # 在Excel中查找该指标
            measured_value = ''
            for row_idx in range(2, detection_ws.max_row + 1):
                cell_indicator = detection_ws.cell(row_idx, indicator_name_col).value
                if cell_indicator and str(cell_indicator).strip() == indicator_name:
                    # 读取检测值（C列，第3列）
                    cell_value = detection_ws.cell(row_idx, value_col).value
                    if cell_value is not None:
                        measured_value = str(cell_value).strip()
                    break

            detection_data.append({
                'indicator_id': indicator_id,
                'indicator_name': indicator_name,
                'unit': unit,
                'measured_value': measured_value or default_value or '',
                'limit_value': limit_value,
                'group_name': group_name
            })

        conn.close()

        # 清理临时文件
        try:
            os.remove(template_path)
            os.remove(detection_path)
        except:
            pass

        # 返回解析结果
        result = {
            'basic_info': basic_info,
            'template_fields': template_fields,
            'detection_data': detection_data
        }

        log_operation('解析报告Excel', f'模板ID: {template_id}, 样品类型ID: {sample_type_id}')

        return jsonify(result), 200

    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print("解析Excel错误详情:")
        print(error_trace)

        # 提供更详细的错误信息
        error_msg = str(e)
        if 'no such table' in error_msg:
            table_name = error_msg.split('no such table:')[-1].strip()
            error_msg = f'数据库表不存在：{table_name}。请检查数据库是否已正确初始化。'
        elif 'no such column' in error_msg:
            error_msg = f'数据库字段不存在：{error_msg}。请检查数据库结构是否正确。'
        elif 'Worksheet' in error_msg and 'does not exist' in error_msg:
            error_msg = f'Excel工作表不存在：{error_msg}。请检查Excel文件格式。'

        return jsonify({
            'error': f'解析失败: {error_msg}',
            'details': error_trace if app.debug else None
        }), 500

@app.route('/api/import-reports', methods=['POST'])
@login_required
def api_import_reports():
    """批量导入报告"""
    from import_processor import import_reports_from_excel

    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400

    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': '仅支持Excel文件(.xlsx, .xls)'}), 400

    # 保存上传的文件
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"import_{timestamp}_{file.filename}"
    upload_path = os.path.join('uploads', filename)
    os.makedirs('uploads', exist_ok=True)
    file.save(upload_path)

    # 获取参数
    template_id = request.form.get('template_id')
    template_id = int(template_id) if template_id else None
    created_by = session.get('username', 'system')

    try:
        # 处理导入
        results = import_reports_from_excel(upload_path, template_id, created_by)

        log_operation('批量导入报告',
                     f'成功:{len(results["success"])} 失败:{len(results["errors"])} 警告:{len(results["warnings"])}')

        return jsonify(results)
    except Exception as e:
        return jsonify({'error': f'导入失败: {str(e)}'}), 500
    finally:
        # 清理临时文件
        if os.path.exists(upload_path):
            os.remove(upload_path)

@app.route('/api/reports/pending-submit', methods=['GET'])
@login_required
def api_reports_pending_submit():
    """获取待提交报告列表（草稿和被拒绝的报告）"""
    conn = get_db_connection()

    # 获取筛选条件
    sample_number = request.args.get('sample_number', '')
    company_id = request.args.get('company_id', '')

    # 构建SQL - 查询当前用户创建的draft和rejected状态的报告
    sql = '''
        SELECT r.*,
               st.name as sample_type_name,
               c.name as company_name,
               t.name as template_name
        FROM reports r
        LEFT JOIN sample_types st ON r.sample_type_id = st.id
        LEFT JOIN companies c ON r.company_id = c.id
        LEFT JOIN excel_report_templates t ON r.template_id = t.id
        WHERE r.created_by = ? AND (r.review_status = 'draft' OR r.review_status = 'rejected' OR r.review_status IS NULL)
    '''
    params = [session['user_id']]

    if sample_number:
        sql += ' AND r.sample_number LIKE ?'
        params.append(f'%{sample_number}%')

    if company_id:
        sql += ' AND r.company_id = ?'
        params.append(company_id)

    sql += ' ORDER BY r.created_at DESC'

    reports = conn.execute(sql, params).fetchall()
    conn.close()

    return jsonify([dict(r) for r in reports])

@app.route('/api/reports/submitted', methods=['GET'])
@login_required
def api_reports_submitted():
    """获取已提交报告列表（pending、approved、rejected状态的报告）"""
    conn = get_db_connection()

    # 获取筛选条件
    sample_number = request.args.get('sample_number', '')
    status = request.args.get('status', '')
    company_id = request.args.get('company_id', '')
    date = request.args.get('date', '')

    # 构建SQL - 查询当前用户创建的已提交报告
    sql = '''
        SELECT r.*,
               st.name as sample_type_name,
               c.name as company_name,
               t.name as template_name
        FROM reports r
        LEFT JOIN sample_types st ON r.sample_type_id = st.id
        LEFT JOIN companies c ON r.company_id = c.id
        LEFT JOIN excel_report_templates t ON r.template_id = t.id
        WHERE r.created_by = ? AND r.review_status IN ('pending', 'approved', 'rejected')
    '''
    params = [session['user_id']]

    if sample_number:
        sql += ' AND r.sample_number LIKE ?'
        params.append(f'%{sample_number}%')

    if status:
        sql += ' AND r.review_status = ?'
        params.append(status)

    if company_id:
        sql += ' AND r.company_id = ?'
        params.append(company_id)

    if date:
        sql += ' AND DATE(r.created_at) = ?'
        params.append(date)

    sql += ' ORDER BY r.created_at DESC'

    reports = conn.execute(sql, params).fetchall()
    conn.close()

    return jsonify([dict(r) for r in reports])

@app.route('/api/reports/review', methods=['GET'])
@login_required
def api_reports_review():
    """获取报告列表（用于审核）"""
    conn = get_db_connection()

    # 获取筛选条件
    status = request.args.get('status', '')
    sample_number = request.args.get('sample_number', '')
    company_id = request.args.get('company_id', '')

    # 构建SQL
    sql = '''
        SELECT r.*,
               st.name as sample_type_name,
               c.name as company_name
        FROM reports r
        LEFT JOIN sample_types st ON r.sample_type_id = st.id
        LEFT JOIN companies c ON r.company_id = c.id
        WHERE 1=1
    '''
    params = []

    if status:
        sql += ' AND r.review_status = ?'
        params.append(status)

    if sample_number:
        sql += ' AND r.sample_number LIKE ?'
        params.append(f'%{sample_number}%')

    if company_id:
        sql += ' AND r.company_id = ?'
        params.append(company_id)

    sql += ' ORDER BY r.created_at DESC'

    reports = conn.execute(sql, params).fetchall()
    conn.close()

    return jsonify([dict(r) for r in reports])

@app.route('/api/reports/<int:id>/review-detail', methods=['GET'])
@login_required
def api_report_review_detail(id):
    """获取报告审核详情"""
    conn = get_db_connection()

    # 获取报告基本信息
    report = conn.execute('''
        SELECT r.*,
               st.name as sample_type_name,
               c.name as company_name,
               t.name as template_name
        FROM reports r
        LEFT JOIN sample_types st ON r.sample_type_id = st.id
        LEFT JOIN companies c ON r.company_id = c.id
        LEFT JOIN excel_report_templates t ON r.template_id = t.id
        WHERE r.id = ?
    ''', (id,)).fetchone()

    if not report:
        conn.close()
        return jsonify({'error': '报告不存在'}), 404

    # 获取检测数据
    detection_data = conn.execute('''
        SELECT rd.*,
               i.name as indicator_name,
               i.unit,
               i.limit_value,
               i.detection_method,
               ig.name as group_name
        FROM report_data rd
        LEFT JOIN indicators i ON rd.indicator_id = i.id
        LEFT JOIN indicator_groups ig ON i.group_id = ig.id
        WHERE rd.report_id = ?
        ORDER BY ig.sort_order, i.sort_order, i.name
    ''', (id,)).fetchall()

    # 获取模板字段值
    template_fields = []
    if report['template_id']:
        template_fields = conn.execute('''
            SELECT rfv.*,
                   tfm.field_name,
                   tfm.field_display_name,
                   tfm.sheet_name,
                   tfm.cell_address
            FROM report_field_values rfv
            LEFT JOIN template_field_mappings tfm ON rfv.field_mapping_id = tfm.id
            WHERE rfv.report_id = ?
        ''', (id,)).fetchall()

    # 获取审核历史记录
    review_history = conn.execute('''
        SELECT rh.*,
               u.username as reviewer_name
        FROM review_history rh
        LEFT JOIN users u ON rh.reviewer_id = u.id
        WHERE rh.report_id = ?
        ORDER BY rh.reviewed_at DESC
    ''', (id,)).fetchall()

    conn.close()

    return jsonify({
        'report': dict(report),
        'detection_data': [dict(d) for d in detection_data],
        'template_fields': [dict(f) for f in template_fields],
        'review_history': [dict(h) for h in review_history]
    })

@app.route('/api/reports/<int:id>/approve', methods=['POST'])
@login_required
def api_approve_report(id):
    """审核通过报告"""
    data = request.json
    comment = data.get('comment', '')

    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # 获取完整报告信息
        report = conn.execute('''
            SELECT r.*, st.name as sample_type_name
            FROM reports r
            LEFT JOIN sample_types st ON r.sample_type_id = st.id
            WHERE r.id = ?
        ''', (id,)).fetchone()

        if not report:
            return jsonify({'error': '报告不存在'}), 404

        review_time = datetime.now()
        username = session.get('username', 'unknown')

        # 更新审核状态
        cursor.execute('''
            UPDATE reports
            SET review_status = 'approved',
                review_person = ?,
                review_time = ?,
                review_comment = ?,
                reviewed_at = ?
            WHERE id = ?
        ''', (username, review_time, comment, review_time, id))

        # 记录审核历史
        cursor.execute('''
            INSERT INTO review_history (report_id, reviewer_id, review_status, review_comment, reviewed_at)
            VALUES (?, ?, 'approved', ?, ?)
        ''', (id, session.get('user_id'), comment, review_time))

        # 自动录入原始数据管理系统
        try:
            # 解析客户信息
            customer_unit = ''
            customer_plant = ''
            if report['remark']:
                try:
                    customer_info = json.loads(report['remark'])
                    customer_unit = customer_info.get('customer_unit', '')
                    customer_plant = customer_info.get('customer_plant', '')
                except:
                    pass

            # 检查是否已存在该样品编号的记录
            existing = cursor.execute(
                'SELECT id FROM raw_data_records WHERE sample_number = ?',
                (report['sample_number'],)
            ).fetchone()

            if existing:
                # 更新现有记录
                cursor.execute('''
                    UPDATE raw_data_records
                    SET report_number = ?,
                        company_name = ?,
                        plant_name = ?,
                        sample_type = ?,
                        sampling_date = ?,
                        updated_at = ?
                    WHERE sample_number = ?
                ''', (report['report_number'], customer_unit, customer_plant, report['sample_type_name'],
                      report['sampling_date'], review_time, report['sample_number']))
                record_id = existing['id']
            else:
                # 插入新记录
                cursor.execute('''
                    INSERT INTO raw_data_records
                    (sample_number, report_number, company_name, plant_name, sample_type, sampling_date, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (report['sample_number'], report['report_number'], customer_unit, customer_plant,
                      report['sample_type_name'], report['sampling_date'], review_time, review_time))
                record_id = cursor.lastrowid

            # 获取检测数据并录入
            detection_data = conn.execute('''
                SELECT rd.*, i.name as indicator_name
                FROM report_data rd
                LEFT JOIN indicators i ON rd.indicator_id = i.id
                WHERE rd.report_id = ?
            ''', (id,)).fetchall()

            # 删除旧的检测值数据
            cursor.execute('DELETE FROM raw_data_values WHERE record_id = ?', (record_id,))

            # 插入新的检测值数据
            for item in detection_data:
                if item['indicator_name'] and item['measured_value']:
                    cursor.execute('''
                        INSERT INTO raw_data_values (record_id, column_name, value)
                        VALUES (?, ?, ?)
                    ''', (record_id, item['indicator_name'], item['measured_value']))

        except Exception as e:
            print(f"自动录入原始数据失败: {e}")
            # 不影响审核流程，继续执行

        conn.commit()
        log_operation('审核报告', f'报告ID: {id}, 结果: 通过', conn=conn)

        return jsonify({'message': '审核通过'})
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/reject', methods=['POST'])
@login_required
def api_reject_report(id):
    """拒绝报告"""
    data = request.json
    comment = data.get('comment', '')

    if not comment:
        return jsonify({'error': '请填写拒绝原因'}), 400

    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # 检查报告是否存在
        report = conn.execute('SELECT id, review_status FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        review_time = datetime.now()
        username = session.get('username', 'unknown')

        # 更新审核状态
        cursor.execute('''
            UPDATE reports
            SET review_status = 'rejected',
                review_person = ?,
                review_time = ?,
                review_comment = ?,
                reviewed_at = ?
            WHERE id = ?
        ''', (username, review_time, comment, review_time, id))

        # 记录审核历史
        cursor.execute('''
            INSERT INTO review_history (report_id, reviewer_id, review_status, review_comment, reviewed_at)
            VALUES (?, ?, 'rejected', ?, ?)
        ''', (id, session.get('user_id'), comment, review_time))

        conn.commit()
        log_operation('审核报告', f'报告ID: {id}, 结果: 拒绝', conn=conn)

        return jsonify({'message': '已拒绝'})
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/submit', methods=['POST'])
@login_required
def api_submit_report(id):
    """提交报告到审核（将draft或rejected状态改为pending）"""
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # 检查报告是否存在
        report = conn.execute('SELECT id, review_status, created_by FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        # 检查权限（仅创建人或管理员可提交）
        if session.get('role') != 'admin' and report['created_by'] != session['user_id']:
            return jsonify({'error': '无权提交此报告'}), 403

        # 检查当前状态是否允许提交
        if report['review_status'] not in ['draft', 'rejected', None]:
            return jsonify({'error': f'当前状态 ({report["review_status"]}) 不允许提交'}), 400

        # 更新状态为pending
        cursor.execute('''
            UPDATE reports
            SET review_status = 'pending'
            WHERE id = ?
        ''', (id,))

        conn.commit()
        log_operation('提交报告', f'报告ID: {id}')

        return jsonify({'message': '报告已提交审核'})
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/return', methods=['POST'])
@login_required
def api_return_report(id):
    """退回报告到审核状态"""
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # 检查报告是否存在
        report = conn.execute('SELECT id, review_status, created_by, review_person FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        # 检查权限（仅管理员或报告创建人可退回）
        if session.get('role') != 'admin' and report['created_by'] != session['user_id']:
            return jsonify({'error': '无权退回此报告'}), 403

        # 检查当前状态是否为已审核
        if report['review_status'] != 'approved':
            return jsonify({'error': f'只有已审核通过的报告才能退回（当前状态: {report["review_status"]}）'}), 400

        # 获取退回原因
        data = request.json or {}
        return_reason = data.get('reason', '').strip()

        # 更新状态为pending，清除生成的报告路径
        cursor.execute('''
            UPDATE reports
            SET review_status = 'pending',
                generated_report_path = NULL,
                review_comment = ?
            WHERE id = ?
        ''', (f'[已退回] {return_reason}' if return_reason else '[已退回] 需要重新审核', id))

        # 记录退回历史
        cursor.execute('''
            INSERT INTO review_history (report_id, reviewer_id, review_status, review_comment, reviewed_at)
            VALUES (?, ?, 'returned', ?, ?)
        ''', (id, session.get('user_id'), return_reason or '退回重新审核', datetime.now()))

        conn.commit()
        log_operation('退回报告', f'报告ID: {id}, 原因: {return_reason}')

        return jsonify({'message': '报告已退回到审核状态'})
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/preview', methods=['POST'])
@login_required
def api_preview_report(id):
    """预览报告数据"""
    from report_generator import ReportGenerator

    data = request.json
    template_id = data.get('template_id')

    if not template_id:
        return jsonify({'error': '请选择报告模板'}), 400

    conn = get_db_connection()

    try:
        # 检查报告是否已审核
        report = conn.execute('''
            SELECT r.*, st.name as sample_type_name
            FROM reports r
            LEFT JOIN sample_types st ON r.sample_type_id = st.id
            WHERE r.id = ?
        ''', (id,)).fetchone()

        if not report:
            return jsonify({'error': '报告不存在'}), 404

        if report['review_status'] != 'approved':
            return jsonify({'error': '只有已审核通过的报告才能预览'}), 400

        # 获取报告数据
        detection_items = conn.execute('''
            SELECT rd.*, i.name, i.unit, i.limit_value, i.detection_method
            FROM report_data rd
            LEFT JOIN indicators i ON rd.indicator_id = i.id
            WHERE rd.report_id = ?
            ORDER BY i.sort_order, i.name
        ''', (id,)).fetchall()

        # 获取模板字段映射
        template_fields = conn.execute('''
            SELECT * FROM template_field_mappings
            WHERE template_id = ?
            ORDER BY sheet_name, cell_address
        ''', (template_id,)).fetchall()

        # 获取模板信息
        template = conn.execute(
            'SELECT * FROM excel_report_templates WHERE id = ?',
            (template_id,)
        ).fetchone()

        # 构建报告数据
        report_data = {
            'report_number': report['report_number'],
            'sample_number': report['sample_number'],
            'sample_type': report['sample_type_name'],
            'detection_date': report['detection_date'],
            'detection_person': report['detection_person'],
            'review_person': report['review_person'],
            'sampling_date': report['sampling_date'],
            'sampler': report['sampler'],
            'sampling_location': report['sampling_location'],
            'report_date': report['report_date'],
        }

        # 尝试从remark中提取客户信息
        if report['remark']:
            try:
                import json
                remark_data = json.loads(report['remark'])
                report_data['customer_unit'] = remark_data.get('customer_unit', '')
                report_data['customer_plant'] = remark_data.get('customer_plant', '')
                # 注意：remark JSON中的键名是 customer_address，需要映射到 unit_address
                report_data['unit_address'] = remark_data.get('customer_address', '') or remark_data.get('unit_address', '')
            except:
                pass

        # 使用ReportGenerator加载完整数据
        generator = ReportGenerator(template_id, report_data, report_id=id)
        generator._load_template_info()
        generator._load_complete_data()

        # 字段名映射表（中文名 -> 英文数据库字段名）
        field_mapping = {
            '报告编号': 'report_number',
            '样品编号': 'sample_number',
            '样品类型': 'sample_type_name',
            '被检单位': 'customer_unit',
            '被检水厂': 'customer_plant',
            '单位地址': 'unit_address',
            '委托单位': 'company_name',
            '采样人': 'sampler',
            '采样日期': 'sampling_date',
            '采样地点': 'sampling_location',
            '采样依据': 'sampling_basis',
            '样品来源': 'sample_source',
            '样品状态': 'sample_status',
            '收样日期': 'sample_received_date',
            '检测日期': 'detection_date',
            '检测人': 'detection_person',
            '检测人员': 'detection_person',
            '审核人': 'review_person',
            '审核人员': 'review_person',
            '报告编制日期': 'report_date',
            '产品标准': 'product_standard',
            '检测项目': 'detection_items_description',
            '检测结论': 'test_conclusion',
            '附加信息': 'additional_info',
            '附件信息': 'attachment_info',
            '备注': 'remark'
        }

        # 打印调试信息
        print(f"\n=== 预览API调试信息 ===")
        print(f"报告ID: {id}")
        print(f"模板ID: {template_id}")
        print(f"generator.report_data 键: {list(generator.report_data.keys())}")
        print(f"模板字段数量: {len(template_fields)}")

        # 构建预览数据
        preview_fields = []
        for field in template_fields:
            field_dict = dict(field)
            field_name = field_dict['field_name']
            is_reference = field_dict.get('is_reference', False)

            # 获取字段值
            if is_reference:
                value = generator._get_reference_value(field_name) if hasattr(generator, '_get_reference_value') else ''
                value_source = '来自已审核报告'
                print(f"引用字段 [{field_name}] = '{value}'")
            else:
                # 将中文字段名映射到英文数据库字段名
                db_field_name = field_mapping.get(field_name, field_name)
                value = generator.report_data.get(db_field_name, field_dict.get('default_value', ''))
                value_source = '当前报告数据'
                print(f"普通字段 [{field_name}] -> [{db_field_name}] = '{value}' (在report_data中: {db_field_name in generator.report_data})")

            preview_fields.append({
                'field_name': field_name,
                'field_display_name': field_dict.get('field_display_name', field_name),
                'sheet_name': field_dict['sheet_name'],
                'cell_address': field_dict['cell_address'],
                'value': value if value is not None else '',
                'value_source': value_source,
                'is_reference': is_reference
            })

        print(f"=== 预览API调试结束 ===\n")

        # 检测数据
        detection_data = [
            {
                'name': item['name'],
                'unit': item['unit'],
                'result': item['measured_value'],
                'limit': item['limit_value'],
                'method': item['detection_method']
            }
            for item in detection_items
        ]

        return jsonify({
            'template_name': template['name'],
            'report_number': report['report_number'],
            'sample_number': report['sample_number'],
            'fields': preview_fields,
            'detection_items': detection_data
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'预览失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/generate', methods=['POST'])
@login_required
def api_generate_report(id):
    """生成最终报告"""
    from report_generator import ReportGenerator

    data = request.json
    template_id = data.get('template_id')
    export_format = data.get('export_format', 'xlsx')  # 导出格式：xlsx 或 pdf
    filename_template = data.get('filename_template')  # 文件名模板（可选）

    if not template_id:
        return jsonify({'error': '请选择报告模板'}), 400

    # 验证导出格式
    if export_format not in ['xlsx', 'pdf']:
        return jsonify({'error': '导出格式必须是 xlsx 或 pdf'}), 400

    conn = get_db_connection()

    try:
        # 检查报告是否已审核
        report = conn.execute('SELECT * FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        if report['review_status'] != 'approved':
            return jsonify({'error': '只有已审核通过的报告才能生成'}), 400

        # 获取报告数据
        detection_items = conn.execute('''
            SELECT rd.*, i.name, i.unit, i.limit_value, i.detection_method
            FROM report_data rd
            LEFT JOIN indicators i ON rd.indicator_id = i.id
            WHERE rd.report_id = ?
        ''', (id,)).fetchall()

        # 构建报告数据
        report_data = {
            'report_number': report['report_number'],
            'sample_number': report['sample_number'],
            'detection_date': report['detection_date'],
            'detection_person': report['detection_person'],
            'review_person': report['review_person'],
            'detection_items': [
                {
                    'name': item['name'],
                    'unit': item['unit'],
                    'result': item['measured_value'],
                    'limit': item['limit_value'],
                    'method': item['detection_method']
                }
                for item in detection_items
            ]
        }

        # 生成报告（传递report_id以从数据库加载完整数据）
        generator = ReportGenerator(template_id, report_data, report_id=id)
        output_path = generator.generate(
            filename_template=filename_template,
            export_format=export_format
        )

        # 更新报告记录
        cursor = conn.cursor()
        cursor.execute(
            'UPDATE reports SET generated_report_path = ? WHERE id = ?',
            (output_path, id)
        )
        conn.commit()

        log_operation('生成报告', f'报告ID: {id}')

        return jsonify({
            'message': '生成成功',
            'file_path': output_path
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'生成失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/download', methods=['GET'])
@login_required
def api_download_report(id):
    """下载生成的报告"""
    conn = get_db_connection()

    report = conn.execute(
        'SELECT generated_report_path FROM reports WHERE id = ?',
        (id,)
    ).fetchone()

    conn.close()

    if not report or not report['generated_report_path']:
        return jsonify({'error': '报告文件不存在'}), 404

    file_path = report['generated_report_path']
    if not os.path.exists(file_path):
        return jsonify({'error': '文件不存在'}), 404

    return send_file(file_path, as_attachment=True, download_name=os.path.basename(file_path))

@app.route('/api/reports/<int:id>/download-pdf', methods=['GET'])
@login_required
def api_download_report_pdf(id):
    """将已生成的报告转换为PDF下载"""
    import subprocess

    conn = get_db_connection()
    report = conn.execute(
        'SELECT generated_report_path, report_number FROM reports WHERE id = ?',
        (id,)
    ).fetchone()
    conn.close()

    if not report or not report['generated_report_path']:
        return jsonify({'error': '报告尚未生成，请先生成报告'}), 404

    file_path = report['generated_report_path']
    if not os.path.exists(file_path):
        return jsonify({'error': '报告文件不存在'}), 404

    # 创建临时副本，设置打印页面参数后再转换
    import shutil
    try:
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S%f')
        temp_xlsx = os.path.join(os.path.dirname(file_path), f'_pdf_temp_{timestamp}.xlsx')
        shutil.copy2(file_path, temp_xlsx)

        # 设置每个工作表的页面布局：缩放至1页宽，自动高度
        wb = openpyxl.load_workbook(temp_xlsx)
        for ws in wb.worksheets:
            ws.page_setup.paperSize = ws.PAPERSIZE_A4
            ws.page_setup.orientation = ws.page_setup.orientation or 'portrait'
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0
            ws.sheet_properties.pageSetUpPr.fitToPage = True
            ws.page_margins.left = 0.5
            ws.page_margins.right = 0.5
            ws.page_margins.top = 0.6
            ws.page_margins.bottom = 0.6
            ws.page_margins.header = 0.3
            ws.page_margins.footer = 0.3
        wb.save(temp_xlsx)
        wb.close()

        abs_path = os.path.abspath(temp_xlsx)
        abs_outdir = os.path.dirname(abs_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', abs_outdir, abs_path],
            capture_output=True, text=True, timeout=60
        )
        pdf_path = abs_path.rsplit('.', 1)[0] + '.pdf'
        # 清理临时Excel
        try:
            os.remove(temp_xlsx)
        except Exception:
            pass
        if result.returncode == 0 and os.path.exists(pdf_path):
            report_number = report['report_number'] or f'report_{id}'
            log_operation('下载PDF报告', f'报告编号:{report_number}')
            response = send_file(pdf_path, as_attachment=True, download_name=f"{report_number}.pdf")
            # 下载后清理临时PDF文件
            @response.call_on_close
            def cleanup():
                try:
                    os.remove(pdf_path)
                except Exception:
                    pass
            return response
        else:
            return jsonify({'error': f'PDF转换失败: {result.stderr}'}), 500
    except subprocess.TimeoutExpired:
        return jsonify({'error': 'PDF转换超时'}), 500
    except Exception as e:
        return jsonify({'error': f'PDF转换异常: {str(e)}'}), 500

@app.route('/api/reports/<int:id>/load-template', methods=['GET'])
@login_required
def api_load_report_template(id):
    """加载报告Excel模板用于在线编辑"""
    from report_generator import ReportGenerator

    conn = get_db_connection()

    try:
        # 获取报告信息
        report = conn.execute('SELECT * FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        # 获取已生成的报告文件路径
        generated_path = report['generated_report_path']

        # 如果还没有生成过报告，需要先获取模板ID
        if not generated_path or not os.path.exists(generated_path):
            # 获取该样品类型的默认模板
            template = conn.execute('''
                SELECT id FROM excel_report_templates
                WHERE sample_type_id = ? AND is_active = 1
                ORDER BY created_at DESC
                LIMIT 1
            ''', (report['sample_type_id'],)).fetchone()

            if not template:
                return jsonify({'error': '未找到对应的报告模板'}), 404

            # 获取报告数据
            detection_items = conn.execute('''
                SELECT rd.*, i.name, i.unit, i.limit_value, i.detection_method
                FROM report_data rd
                LEFT JOIN indicators i ON rd.indicator_id = i.id
                WHERE rd.report_id = ?
            ''', (id,)).fetchall()

            # 构建报告数据
            report_data = {
                'report_number': report['report_number'],
                'sample_number': report['sample_number'],
                'detection_date': report['detection_date'],
                'detection_person': report['detection_person'],
                'review_person': report['review_person'],
                'detection_items': [
                    {
                        'name': item['name'],
                        'unit': item['unit'],
                        'result': item['measured_value'],
                        'limit': item['limit_value'],
                        'method': item['detection_method']
                    }
                    for item in detection_items
                ]
            }

            # 生成临时报告
            generator = ReportGenerator(template['id'], report_data, report_id=id)
            generated_path = generator.generate()

        # 读取Excel文件并转换为x-spreadsheet格式
        wb = openpyxl.load_workbook(generated_path, data_only=False)
        sheets_data = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_data = {
                'name': sheet_name,
                'rows': {},
                'cols': {}
            }

            # 读取所有单元格数据
            for row_idx, row in enumerate(ws.iter_rows()):
                if row_idx >= 1000:  # 限制最多读取1000行
                    break

                cells = {}
                for col_idx, cell in enumerate(row):
                    if col_idx >= 50:  # 限制最多读取50列
                        break

                    if cell.value is not None:
                        cells[col_idx] = {
                            'text': str(cell.value)
                        }

                if cells:
                    sheet_data['rows'][row_idx] = {'cells': cells}

            sheets_data.append(sheet_data)

        wb.close()

        return jsonify({
            'sheets': sheets_data,
            'report_id': id,
            'report_number': report['report_number']
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'加载模板失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/reports/<int:id>/save-from-editor', methods=['POST'])
@login_required
def api_save_from_editor(id):
    """从在线编辑器保存并生成新报告"""
    conn = get_db_connection()

    try:
        data = request.json
        sheets_data = data.get('sheets', [])

        if not sheets_data:
            return jsonify({'error': '没有数据'}), 400

        # 获取报告信息
        report = conn.execute('SELECT * FROM reports WHERE id = ?', (id,)).fetchone()
        if not report:
            return jsonify({'error': '报告不存在'}), 404

        # 创建新的Excel文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"exports/report_{report['report_number']}_edited_{timestamp}.xlsx"
        os.makedirs('exports', exist_ok=True)

        # 创建工作簿
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # 删除默认sheet

        # 根据编辑器数据创建sheets
        for sheet_info in sheets_data:
            sheet_name = sheet_info.get('name', 'Sheet1')
            ws = wb.create_sheet(title=sheet_name)

            rows_data = sheet_info.get('rows', {})
            for row_idx_str, row_info in rows_data.items():
                row_idx = int(row_idx_str)
                cells = row_info.get('cells', {})

                for col_idx_str, cell_info in cells.items():
                    col_idx = int(col_idx_str)
                    text = cell_info.get('text', '')

                    # Excel是1-based索引
                    ws.cell(row=row_idx + 1, column=col_idx + 1, value=text)

        # 保存文件
        wb.save(output_path)
        wb.close()

        # 更新报告记录
        cursor = conn.cursor()
        cursor.execute(
            'UPDATE reports SET generated_report_path = ? WHERE id = ?',
            (output_path, id)
        )
        conn.commit()

        log_operation('在线编辑报告', f'报告ID: {id}', conn=conn)

        return jsonify({
            'message': '保存成功',
            'file_path': output_path
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'保存失败: {str(e)}'}), 500
    finally:
        conn.close()

@app.route('/api/template-fields/batch-update-defaults', methods=['POST'])
@admin_required
def api_batch_update_field_defaults():
    """批量更新字段默认值"""
    data = request.json
    updates = data.get('updates', [])

    if not updates:
        return jsonify({'error': '没有要更新的数据'}), 400

    conn = get_db_connection()

    try:
        for update in updates:
            field_id = update.get('id')
            default_value = update.get('default_value', '')

            conn.execute(
                'UPDATE template_field_mappings SET default_value = ? WHERE id = ?',
                (default_value, field_id)
            )

        conn.commit()
        log_operation('更新字段默认值', f'批量更新 {len(updates)} 个字段', conn=conn)
        conn.close()

        return jsonify({'message': f'成功更新 {len(updates)} 个字段的默认值'})
    except Exception as e:
        conn.close()
        return jsonify({'error': f'更新失败: {str(e)}'}), 500

# 此路由已在第2172行定义，此处删除重复定义

# ==================== 原始数据管理 API ====================
from raw_data_importer import RawDataImporter
from raw_data_template_generator import generate_raw_data_template
from werkzeug.utils import secure_filename

UPLOAD_FOLDER = 'temp/uploads'
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/raw-data/upload', methods=['POST'])
@login_required
def api_raw_data_upload():
    """上传并导入Excel原始数据"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '未选择文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': '文件格式不支持，仅支持.xlsx和.xls格式'}), 400

        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        saved_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, saved_filename)
        file.save(filepath)

        # 获取处理选项
        on_duplicate = request.form.get('on_duplicate', 'skip')  # skip, overwrite, abort

        # 导入数据
        importer = RawDataImporter()
        result = importer.import_excel(filepath, on_duplicate=on_duplicate)

        # 删除临时文件
        try:
            os.remove(filepath)
        except:
            pass

        # 记录操作日志
        if result['success']:
            log_operation(
                '导入原始数据',
                f"导入成功: {result['success_count']}条，跳过: {result['skip_count']}条"
            )

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': f'上传失败: {str(e)}'}), 500

@app.route('/api/raw-data/columns', methods=['GET'])
@login_required
def api_raw_data_columns():
    """获取当前系统的列名配置"""
    try:
        importer = RawDataImporter()
        columns = importer.get_column_list()

        if columns is None:
            return jsonify({'columns': None, 'message': '系统尚未初始化，请先导入数据'})

        return jsonify({'columns': columns})

    except Exception as e:
        return jsonify({'error': f'获取列名失败: {str(e)}'}), 500

@app.route('/api/raw-data/download-template', methods=['GET'])
@login_required
def api_raw_data_download_template():
    """下载原始数据导入模板"""
    try:
        # 生成模板
        template_path = generate_raw_data_template()

        # 生成下载文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        download_name = f'原始数据导入模板_{timestamp}.xlsx'

        # 记录操作日志
        log_operation('下载导入模板', '下载原始数据导入模板')

        # 发送文件
        return send_file(
            template_path,
            as_attachment=True,
            download_name=download_name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        return jsonify({'error': f'生成模板失败: {str(e)}'}), 500

@app.route('/api/raw-data/search', methods=['POST'])
@login_required
def api_raw_data_search():
    """根据样品编号精确查询原始数据"""
    try:
        data = request.json
        sample_number = data.get('sample_number', '').strip()

        if not sample_number:
            return jsonify({'error': '样品编号不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 查询主记录
        cursor.execute('''
            SELECT id, sample_number, report_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE sample_number = ?
        ''', (sample_number,))

        record = cursor.fetchone()

        if not record:
            conn.close()
            return jsonify({'found': False, 'message': '未找到该样品编号的数据'})

        record_id = record[0]
        record_data = {
            'id': record[0],
            'sample_number': record[1],
            'report_number': record[2],
            'company_name': record[3],
            'plant_name': record[4],
            'sample_type': record[5],
            'sampling_date': record[6],
            'created_at': record[7],
            'updated_at': record[8]
        }

        # 查询检测指标数据
        cursor.execute('''
            SELECT column_name, value
            FROM raw_data_values
            WHERE record_id = ?
            ORDER BY id
        ''', (record_id,))

        indicators = {}
        for row in cursor.fetchall():
            indicators[row[0]] = row[1]

        conn.close()

        return jsonify({
            'found': True,
            'data': record_data,
            'indicators': indicators
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/search-by-company', methods=['POST'])
@login_required
def api_raw_data_search_by_company():
    """根据被检单位模糊查询原始数据列表"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()

        if not company_name:
            return jsonify({'error': '被检单位不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 模糊查询主记录
        cursor.execute('''
            SELECT id, sample_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE company_name LIKE ?
            ORDER BY company_name, plant_name, sampling_date DESC
        ''', (f'%{company_name}%',))

        records = cursor.fetchall()
        conn.close()

        if not records:
            return jsonify({'found': False, 'message': '未找到匹配的数据', 'records': []})

        result_list = []
        for record in records:
            result_list.append({
                'id': record[0],
                'sample_number': record[1],
                'company_name': record[2],
                'plant_name': record[3],
                'sample_type': record[4],
                'sampling_date': record[5],
                'created_at': record[6],
                'updated_at': record[7]
            })

        return jsonify({
            'found': True,
            'count': len(result_list),
            'records': result_list
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/search-by-plant', methods=['POST'])
@login_required
def api_raw_data_search_by_plant():
    """根据被检水厂模糊查询原始数据列表"""
    try:
        data = request.json
        plant_name = data.get('plant_name', '').strip()

        if not plant_name:
            return jsonify({'error': '被检水厂不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 模糊查询主记录
        cursor.execute('''
            SELECT id, sample_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE plant_name LIKE ?
            ORDER BY company_name, plant_name, sampling_date DESC
        ''', (f'%{plant_name}%',))

        records = cursor.fetchall()
        conn.close()

        if not records:
            return jsonify({'found': False, 'message': '未找到匹配的数据', 'records': []})

        result_list = []
        for record in records:
            result_list.append({
                'id': record[0],
                'sample_number': record[1],
                'company_name': record[2],
                'plant_name': record[3],
                'sample_type': record[4],
                'sampling_date': record[5],
                'created_at': record[6],
                'updated_at': record[7]
            })

        return jsonify({
            'found': True,
            'count': len(result_list),
            'records': result_list
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/search-companies', methods=['POST'])
@login_required
def api_raw_data_search_companies():
    """根据关键词模糊查找被检单位列表"""
    try:
        data = request.json
        keyword = data.get('keyword', '').strip()

        if not keyword:
            return jsonify({'error': '搜索关键词不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 模糊查询所有匹配的单位
        cursor.execute('''
            SELECT DISTINCT company_name
            FROM raw_data_records
            WHERE company_name LIKE ?
            ORDER BY company_name
        ''', (f'%{keyword}%',))

        companies = [row[0] for row in cursor.fetchall() if row[0]]
        conn.close()

        return jsonify({
            'companies': companies,
            'count': len(companies)
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/get-plants', methods=['POST'])
@login_required
def api_raw_data_get_plants():
    """根据被检单位获取水厂列表"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()

        if not company_name:
            return jsonify({'error': '被检单位不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 查询该单位下的所有水厂
        cursor.execute('''
            SELECT DISTINCT plant_name
            FROM raw_data_records
            WHERE company_name = ?
            ORDER BY plant_name
        ''', (company_name,))

        plants = [row[0] for row in cursor.fetchall() if row[0]]
        conn.close()

        return jsonify({
            'plants': plants,
            'count': len(plants)
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/get-sample-types', methods=['POST'])
@login_required
def api_raw_data_get_sample_types():
    """根据被检单位和水厂获取样品类型列表"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()
        plant_names = data.get('plant_names', [])  # 可以是多个水厂

        if not company_name:
            return jsonify({'error': '被检单位不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        if plant_names:
            # 查询指定水厂的样品类型
            placeholders = ','.join(['?' for _ in plant_names])
            query = f'''
                SELECT DISTINCT sample_type
                FROM raw_data_records
                WHERE company_name = ? AND plant_name IN ({placeholders})
                ORDER BY sample_type
            '''
            cursor.execute(query, [company_name] + plant_names)
        else:
            # 查询该单位下所有样品类型
            cursor.execute('''
                SELECT DISTINCT sample_type
                FROM raw_data_records
                WHERE company_name = ?
                ORDER BY sample_type
            ''', (company_name,))

        sample_types = [row[0] for row in cursor.fetchall() if row[0]]
        conn.close()

        return jsonify({
            'sample_types': sample_types,
            'count': len(sample_types)
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/search-by-filters', methods=['POST'])
@login_required
def api_raw_data_search_by_filters():
    """根据单位、水厂、样品类型组合查询（支持录入时间筛选）"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()
        plant_names = data.get('plant_names', [])
        sample_types = data.get('sample_types', [])
        created_start = data.get('created_start', '').strip()  # 录入开始日期
        created_end = data.get('created_end', '').strip()      # 录入结束日期

        if not company_name:
            return jsonify({'error': '被检单位不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 构建查询条件
        conditions = ['company_name = ?']
        params = [company_name]

        if plant_names:
            placeholders = ','.join(['?' for _ in plant_names])
            conditions.append(f'plant_name IN ({placeholders})')
            params.extend(plant_names)

        if sample_types:
            placeholders = ','.join(['?' for _ in sample_types])
            conditions.append(f'sample_type IN ({placeholders})')
            params.extend(sample_types)

        # 添加录入时间筛选条件
        if created_start:
            conditions.append('DATE(created_at) >= ?')
            params.append(created_start)

        if created_end:
            conditions.append('DATE(created_at) <= ?')
            params.append(created_end)

        where_clause = ' AND '.join(conditions)

        query = f'''
            SELECT id, sample_number, report_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE {where_clause}
            ORDER BY created_at DESC, company_name, plant_name, sampling_date DESC
        '''

        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()

        if not records:
            return jsonify({'found': False, 'message': '未找到匹配的数据', 'records': []})

        result_list = []
        for record in records:
            result_list.append({
                'id': record[0],
                'sample_number': record[1],
                'report_number': record[2],
                'company_name': record[3],
                'plant_name': record[4],
                'sample_type': record[5],
                'sampling_date': record[6],
                'created_at': record[7],
                'updated_at': record[8]
            })

        return jsonify({
            'found': True,
            'count': len(result_list),
            'records': result_list
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/search-by-time', methods=['POST'])
@login_required
def api_raw_data_search_by_time():
    """按录入时间查询所有样品数据"""
    try:
        data = request.json
        created_start = data.get('created_start', '').strip()
        created_end = data.get('created_end', '').strip()

        if not created_start and not created_end:
            return jsonify({'error': '请至少选择开始日期或结束日期'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 构建查询条件
        conditions = []
        params = []

        if created_start:
            conditions.append('DATE(created_at) >= ?')
            params.append(created_start)

        if created_end:
            conditions.append('DATE(created_at) <= ?')
            params.append(created_end)

        where_clause = ' AND '.join(conditions)

        query = f'''
            SELECT id, sample_number, report_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE {where_clause}
            ORDER BY created_at DESC, sample_number
        '''

        cursor.execute(query, params)
        records = cursor.fetchall()
        conn.close()

        if not records:
            return jsonify({'found': False, 'message': '未找到匹配的数据', 'records': []})

        result_list = []
        for record in records:
            result_list.append({
                'id': record[0],
                'sample_number': record[1],
                'report_number': record[2],
                'company_name': record[3],
                'plant_name': record[4],
                'sample_type': record[5],
                'sampling_date': record[6],
                'created_at': record[7],
                'updated_at': record[8]
            })

        return jsonify({
            'found': True,
            'count': len(result_list),
            'records': result_list
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

@app.route('/api/raw-data/detail/<int:record_id>', methods=['GET'])
@login_required
def api_raw_data_detail(record_id):
    """获取原始数据详情"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # 查询主记录
        cursor.execute('''
            SELECT id, sample_number, report_number, company_name, plant_name, sample_type, sampling_date,
                   created_at, updated_at
            FROM raw_data_records
            WHERE id = ?
        ''', (record_id,))

        record = cursor.fetchone()

        if not record:
            conn.close()
            return jsonify({'error': '记录不存在'}), 404

        record_data = {
            'id': record[0],
            'sample_number': record[1],
            'report_number': record[2],
            'company_name': record[3],
            'plant_name': record[4],
            'sample_type': record[5],
            'sampling_date': record[6],
            'created_at': record[7],
            'updated_at': record[8]
        }

        # 查询检测指标数据
        cursor.execute('''
            SELECT column_name, value
            FROM raw_data_values
            WHERE record_id = ?
            ORDER BY id
        ''', (record_id,))

        indicators = {}
        for row in cursor.fetchall():
            indicators[row[0]] = row[1]

        conn.close()

        return jsonify({
            'data': record_data,
            'indicators': indicators
        })

    except Exception as e:
        return jsonify({'error': f'获取详情失败: {str(e)}'}), 500

@app.route('/api/raw-data/update/<int:record_id>', methods=['PUT'])
@login_required
def api_raw_data_update(record_id):
    """更新原始数据记录"""
    try:
        data = request.json

        # 提取基础字段
        sample_number = data.get('sample_number', '').strip()
        company_name = data.get('company_name', '').strip()
        plant_name = data.get('plant_name', '').strip()
        sample_type = data.get('sample_type', '').strip()
        sampling_date = data.get('sampling_date', '').strip()
        indicators = data.get('indicators', {})

        if not sample_number:
            return jsonify({'error': '样品编号不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 检查记录是否存在
        cursor.execute('SELECT id FROM raw_data_records WHERE id = ?', (record_id,))
        if not cursor.fetchone():
            conn.close()
            return jsonify({'error': '记录不存在'}), 404

        # 检查样品编号是否与其他记录重复
        cursor.execute('SELECT id FROM raw_data_records WHERE sample_number = ? AND id != ?',
                      (sample_number, record_id))
        if cursor.fetchone():
            conn.close()
            return jsonify({'error': f'样品编号"{sample_number}"已被其他记录使用'}), 400

        # 更新主记录
        cursor.execute('''
            UPDATE raw_data_records
            SET sample_number = ?, company_name = ?, plant_name = ?,
                sample_type = ?, sampling_date = ?, updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        ''', (sample_number, company_name, plant_name, sample_type, sampling_date, record_id))

        # 删除旧的检测值数据
        cursor.execute('DELETE FROM raw_data_values WHERE record_id = ?', (record_id,))

        # 插入新的检测值数据
        for column_name, value in indicators.items():
            if value is not None and str(value).strip():
                cursor.execute('''
                    INSERT INTO raw_data_values (record_id, column_name, value)
                    VALUES (?, ?, ?)
                ''', (record_id, column_name, str(value).strip()))

        conn.commit()
        conn.close()

        return jsonify({'message': '更新成功'})

    except Exception as e:
        if conn:
            conn.rollback()
            conn.close()
        return jsonify({'error': f'更新失败: {str(e)}'}), 500

@app.route('/api/raw-data/delete/<int:record_id>', methods=['DELETE'])
@login_required
def api_raw_data_delete(record_id):
    """删除原始数据记录"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # 检查记录是否存在
        cursor.execute('SELECT sample_number FROM raw_data_records WHERE id = ?', (record_id,))
        record = cursor.fetchone()

        if not record:
            conn.close()
            return jsonify({'error': '记录不存在'}), 404

        # 删除记录（级联删除会自动删除关联的检测值）
        cursor.execute('DELETE FROM raw_data_records WHERE id = ?', (record_id,))

        conn.commit()
        conn.close()

        return jsonify({'message': f'已删除样品编号"{record[0]}"的记录'})

    except Exception as e:
        if conn:
            conn.rollback()
            conn.close()
        return jsonify({'error': f'删除失败: {str(e)}'}), 500

@app.route('/api/raw-data/export-single', methods=['POST'])
@login_required
def api_raw_data_export_single():
    """导出单条原始数据记录"""
    try:
        data = request.json
        sample_number = data.get('sample_number', '').strip()
        export_format = data.get('format', 'excel')  # excel 或 csv

        if not sample_number:
            return jsonify({'error': '样品编号不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 查询数据
        cursor.execute('''
            SELECT id, sample_number, company_name, plant_name, sample_type, sampling_date
            FROM raw_data_records
            WHERE sample_number = ?
        ''', (sample_number,))

        record = cursor.fetchone()

        if not record:
            conn.close()
            return jsonify({'error': '未找到该样品编号的数据'}), 404

        record_id = record[0]

        # 获取列名顺序
        cursor.execute('''
            SELECT column_name
            FROM raw_data_column_schema
            ORDER BY column_order
        ''')
        columns = [row[0] for row in cursor.fetchall()]

        # 构建数据行
        data_row = {
            '样品编号': record[1],
            '所属公司': record[2],
            '所属水厂': record[3],
            '水样类型': record[4],
            '采样时间': record[5]
        }

        # 获取检测指标值
        cursor.execute('''
            SELECT column_name, value
            FROM raw_data_values
            WHERE record_id = ?
        ''', (record_id,))

        for row in cursor.fetchall():
            data_row[row[0]] = row[1]

        conn.close()

        # 创建DataFrame
        df = pd.DataFrame([data_row], columns=columns)

        # 生成文件
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        if export_format == 'csv':
            filename = f'{sample_number}_{timestamp}.csv'
            filepath = os.path.join('exports', filename)
            df.to_csv(filepath, index=False, encoding='utf-8-sig')
        else:
            filename = f'{sample_number}_{timestamp}.xlsx'
            filepath = os.path.join('exports', filename)
            df.to_excel(filepath, index=False, engine='openpyxl')

        log_operation('导出单条原始数据', f'样品编号: {sample_number}')

        return send_file(filepath, as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

# ==================== 导出模板管理 API ====================

@app.route('/api/export-templates/categories', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_export_template_categories():
    """导出模板分类管理"""
    conn = get_db_connection()

    try:
        if request.method == 'GET':
            # 获取所有分类
            cursor = conn.cursor()
            cursor.execute('''
                SELECT id, name, sort_order, created_at
                FROM export_template_categories
                ORDER BY sort_order, id
            ''')
            categories = []
            for row in cursor.fetchall():
                categories.append({
                    'id': row[0],
                    'name': row[1],
                    'sort_order': row[2],
                    'created_at': row[3]
                })
            conn.close()
            return jsonify({'categories': categories})

        elif request.method == 'POST':
            # 创建新分类
            data = request.json
            name = data.get('name', '').strip()

            if not name:
                conn.close()
                return jsonify({'error': '分类名称不能为空'}), 400

            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO export_template_categories (name, sort_order)
                VALUES (?, ?)
            ''', (name, data.get('sort_order', 0)))

            category_id = cursor.lastrowid
            conn.commit()
            conn.close()

            log_operation('创建导出模板分类', f'分类名称: {name}')

            return jsonify({'message': '分类创建成功', 'category_id': category_id})

        elif request.method == 'DELETE':
            # 删除分类
            data = request.json
            category_id = data.get('category_id')

            if not category_id:
                conn.close()
                return jsonify({'error': '分类ID不能为空'}), 400

            cursor = conn.cursor()
            cursor.execute('DELETE FROM export_template_categories WHERE id = ?', (category_id,))
            conn.commit()
            conn.close()

            log_operation('删除导出模板分类', f'分类ID: {category_id}')

            return jsonify({'message': '分类删除成功'})

    except Exception as e:
        conn.close()
        return jsonify({'error': f'操作失败: {str(e)}'}), 500

@app.route('/api/export-templates', methods=['GET', 'POST'])
@login_required
def api_export_templates_list():
    """获取或创建导出模板"""
    conn = get_db_connection()

    try:
        if request.method == 'GET':
            # 获取所有模板
            cursor = conn.cursor()
            cursor.execute('''
                SELECT et.id, et.category_id, etc.name as category_name,
                       et.sample_type_id, st.name as sample_type_name,
                       et.name, et.description, et.created_at, et.updated_at
                FROM export_templates et
                LEFT JOIN export_template_categories etc ON et.category_id = etc.id
                LEFT JOIN sample_types st ON et.sample_type_id = st.id
                ORDER BY etc.sort_order, et.id
            ''')

            templates = []
            for row in cursor.fetchall():
                template_id = row[0]

                # 获取模板包含的列
                cursor.execute('''
                    SELECT column_name, column_order
                    FROM export_template_columns
                    WHERE template_id = ?
                    ORDER BY column_order
                ''', (template_id,))

                columns = [col[0] for col in cursor.fetchall()]

                templates.append({
                    'id': row[0],
                    'category_id': row[1],
                    'category_name': row[2],
                    'sample_type_id': row[3],
                    'sample_type_name': row[4],
                    'name': row[5],
                    'description': row[6],
                    'columns': columns,
                    'created_at': row[7],
                    'updated_at': row[8]
                })

            conn.close()
            return jsonify({'templates': templates})

        elif request.method == 'POST':
            # 创建新模板
            data = request.json
            category_id = data.get('category_id')
            sample_type_id = data.get('sample_type_id')
            name = data.get('name', '').strip()
            description = data.get('description', '').strip()
            columns = data.get('columns', [])

            if not name:
                conn.close()
                return jsonify({'error': '模板名称不能为空'}), 400

            if not sample_type_id:
                conn.close()
                return jsonify({'error': '请选择样品类型'}), 400

            if not columns:
                conn.close()
                return jsonify({'error': '至少选择一个检测指标'}), 400

            cursor = conn.cursor()

            # 插入模板
            cursor.execute('''
                INSERT INTO export_templates (category_id, sample_type_id, name, description)
                VALUES (?, ?, ?, ?)
            ''', (category_id, sample_type_id, name, description))

            template_id = cursor.lastrowid

            # 插入模板列
            for idx, col_name in enumerate(columns):
                cursor.execute('''
                    INSERT INTO export_template_columns (template_id, column_name, column_order)
                    VALUES (?, ?, ?)
                ''', (template_id, col_name, idx))

            conn.commit()
            conn.close()

            log_operation('创建导出模板', f'模板名称: {name}，样品类型ID: {sample_type_id}，包含{len(columns)}个指标')

            return jsonify({'message': '模板创建成功', 'template_id': template_id})

    except Exception as e:
        conn.close()
        return jsonify({'error': f'操作失败: {str(e)}'}), 500

@app.route('/api/export-templates/<int:template_id>', methods=['PUT', 'DELETE'])
@login_required
def api_export_template_detail(template_id):
    """修改或删除导出模板"""
    conn = get_db_connection()

    try:
        if request.method == 'PUT':
            # 修改模板
            data = request.json
            sample_type_id = data.get('sample_type_id')
            name = data.get('name', '').strip()
            description = data.get('description', '').strip()
            columns = data.get('columns', [])

            if not name:
                conn.close()
                return jsonify({'error': '模板名称不能为空'}), 400

            if not sample_type_id:
                conn.close()
                return jsonify({'error': '请选择样品类型'}), 400

            cursor = conn.cursor()

            # 更新模板基本信息
            cursor.execute('''
                UPDATE export_templates
                SET sample_type_id = ?, name = ?, description = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (sample_type_id, name, description, template_id))

            # 删除旧的列配置
            cursor.execute('DELETE FROM export_template_columns WHERE template_id = ?', (template_id,))

            # 插入新的列配置
            for idx, col_name in enumerate(columns):
                cursor.execute('''
                    INSERT INTO export_template_columns (template_id, column_name, column_order)
                    VALUES (?, ?, ?)
                ''', (template_id, col_name, idx))

            conn.commit()
            conn.close()

            log_operation('修改导出模板', f'模板ID: {template_id}')

            return jsonify({'message': '模板修改成功'})

        elif request.method == 'DELETE':
            # 删除模板
            cursor = conn.cursor()
            cursor.execute('DELETE FROM export_templates WHERE id = ?', (template_id,))
            conn.commit()
            conn.close()

            log_operation('删除导出模板', f'模板ID: {template_id}')

            return jsonify({'message': '模板删除成功'})

    except Exception as e:
        conn.close()
        return jsonify({'error': f'操作失败: {str(e)}'}), 500

@app.route('/api/sample-types/<int:sample_type_id>/indicators', methods=['GET'])
@login_required
def api_sample_type_indicators(sample_type_id):
    """获取指定样品类型的所有检测指标"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # 获取该样品类型关联的检测指标
        cursor.execute('''
            SELECT i.id, i.name, i.unit, ti.is_required, ti.sort_order
            FROM template_indicators ti
            JOIN indicators i ON ti.indicator_id = i.id
            WHERE ti.sample_type_id = ?
            ORDER BY ti.sort_order, i.name
        ''', (sample_type_id,))

        indicators = []
        for row in cursor.fetchall():
            indicators.append({
                'id': row[0],
                'name': row[1],
                'unit': row[2],
                'is_required': bool(row[3]),
                'sort_order': row[4]
            })

        conn.close()
        return jsonify({'indicators': indicators})

    except Exception as e:
        return jsonify({'error': f'获取指标失败: {str(e)}'}), 500

@app.route('/api/raw-data/companies', methods=['GET'])
@login_required
def api_raw_data_companies():
    """获取所有被检单位"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute('''
            SELECT DISTINCT company_name
            FROM raw_data_records
            WHERE company_name IS NOT NULL AND company_name != ''
            ORDER BY company_name
        ''')

        companies = [row[0] for row in cursor.fetchall()]
        conn.close()

        return jsonify({'companies': companies})

    except Exception as e:
        return jsonify({'error': f'获取被检单位失败: {str(e)}'}), 500

@app.route('/api/raw-data/plants', methods=['POST'])
@login_required
def api_raw_data_plants():
    """获取指定被检单位的所有水厂"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()

        if not company_name:
            return jsonify({'error': '被检单位不能为空'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute('''
            SELECT DISTINCT plant_name
            FROM raw_data_records
            WHERE company_name = ? AND plant_name IS NOT NULL AND plant_name != ''
            ORDER BY plant_name
        ''', (company_name,))

        plants = [row[0] for row in cursor.fetchall()]
        conn.close()

        return jsonify({'plants': plants})

    except Exception as e:
        return jsonify({'error': f'获取水厂失败: {str(e)}'}), 500

@app.route('/api/raw-data/samples', methods=['POST'])
@login_required
def api_raw_data_samples():
    """获取指定条件下的所有样品"""
    try:
        data = request.json
        company_name = data.get('company_name', '').strip()
        plant_name = data.get('plant_name', '').strip()
        sample_type_id = data.get('sample_type_id')

        conn = get_db_connection()
        cursor = conn.cursor()

        # 构建查询
        query = '''
            SELECT DISTINCT rdr.id, rdr.sample_number, rdr.report_number,
                   rdr.sampling_date, st.name as sample_type_name
            FROM raw_data_records rdr
            LEFT JOIN sample_types st ON rdr.sample_type = st.name
            WHERE 1=1
        '''
        params = []

        if company_name:
            query += ' AND rdr.company_name = ?'
            params.append(company_name)

        if plant_name:
            query += ' AND rdr.plant_name = ?'
            params.append(plant_name)

        if sample_type_id:
            # 根据sample_type_id获取样品类型名称
            cursor.execute('SELECT name FROM sample_types WHERE id = ?', (sample_type_id,))
            result = cursor.fetchone()
            if result:
                query += ' AND rdr.sample_type = ?'
                params.append(result[0])

        query += ' ORDER BY rdr.sampling_date DESC, rdr.sample_number'

        cursor.execute(query, params)

        samples = []
        for row in cursor.fetchall():
            samples.append({
                'id': row[0],
                'sample_number': row[1],
                'report_number': row[2],
                'sampling_date': row[3],
                'sample_type_name': row[4]
            })

        conn.close()
        return jsonify({'samples': samples})

    except Exception as e:
        return jsonify({'error': f'获取样品失败: {str(e)}'}), 500

@app.route('/api/raw-data/filter-export', methods=['POST'])
@login_required
def api_raw_data_filter_export():
    """筛选并导出原始数据 - 新版：行为样品编号，列为检测项目"""
    try:
        data = request.json
        template_id = data.get('template_id')
        selected_sample_ids = data.get('selected_sample_ids', [])  # 选中的样品ID列表

        if not template_id:
            return jsonify({'error': '请选择导出模板'}), 400

        if not selected_sample_ids:
            return jsonify({'error': '请至少选择一个样品'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        # 获取模板配置的检测指标
        cursor.execute('''
            SELECT column_name
            FROM export_template_columns
            WHERE template_id = ?
            ORDER BY column_order
        ''', (template_id,))

        template_indicators = [row[0] for row in cursor.fetchall()]

        if not template_indicators:
            conn.close()
            return jsonify({'error': '模板配置错误：未包含任何检测指标'}), 400

        # 查询选中的样品记录
        placeholders = ','.join(['?'] * len(selected_sample_ids))
        query = f'''
            SELECT id, sample_number, report_number, company_name, plant_name,
                   sample_type, sampling_date
            FROM raw_data_records
            WHERE id IN ({placeholders})
            ORDER BY sampling_date, sample_number
        '''

        cursor.execute(query, selected_sample_ids)
        records = cursor.fetchall()

        if not records:
            conn.close()
            return jsonify({'error': '未找到选中的样品数据'}), 404

        # 准备导出数据：行为样品编号，列为检测项目
        export_data = []

        for record in records:
            record_id = record[0]
            row_data = {
                '样品编号': record[1],
                '报告编号': record[2] or '',
                '被检单位': record[3] or '',
                '被检水厂': record[4] or '',
                '样品类型': record[5] or '',
                '采样日期': record[6]
            }

            # 获取该样品的所有检测指标值
            cursor.execute('''
                SELECT column_name, value
                FROM raw_data_values
                WHERE record_id = ?
            ''', (record_id,))

            indicator_values = dict(cursor.fetchall())

            # 只导出模板中选择的检测指标
            for indicator in template_indicators:
                row_data[indicator] = indicator_values.get(indicator, '')

            export_data.append(row_data)

        conn.close()

        # 构建列顺序：基础字段 + 检测指标
        columns = ['样品编号', '报告编号', '被检单位', '被检水厂', '样品类型', '采样日期'] + template_indicators

        # 创建DataFrame
        df = pd.DataFrame(export_data, columns=columns)

        # 生成文件
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        filename = f'导出数据_{timestamp}.xlsx'
        filepath = os.path.join('exports', filename)
        df.to_excel(filepath, index=False, engine='openpyxl')

        log_operation('筛选导出原始数据', f'导出{len(records)}条记录，包含{len(template_indicators)}个检测指标')

        return send_file(filepath, as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/api/raw-data/filter-preview', methods=['POST'])
@login_required
def api_raw_data_filter_preview():
    """预览筛选结果（返回符合条件的样品编号列表）"""
    try:
        data = request.json
        filter_field = data.get('filter_field')
        filter_value = data.get('filter_value', '').strip()
        date_start = data.get('date_start')
        date_end = data.get('date_end')

        conn = get_db_connection()
        cursor = conn.cursor()

        # 构建查询
        query = 'SELECT sample_number, company_name, plant_name, sample_type, sampling_date FROM raw_data_records WHERE 1=1'
        params = []

        if filter_field == 'company_name' and filter_value:
            query += ' AND company_name LIKE ?'
            params.append(f'%{filter_value}%')
        elif filter_field == 'plant_name' and filter_value:
            query += ' AND plant_name LIKE ?'
            params.append(f'%{filter_value}%')
        elif filter_field == 'sample_type' and filter_value:
            query += ' AND sample_type LIKE ?'
            params.append(f'%{filter_value}%')
        elif filter_field == 'date_range' and date_start and date_end:
            query += ' AND sampling_date BETWEEN ? AND ?'
            params.append(date_start)
            params.append(date_end)

        query += ' ORDER BY sampling_date DESC'

        cursor.execute(query, params)
        results = []

        for row in cursor.fetchall():
            results.append({
                'sample_number': row[0],
                'company_name': row[1],
                'plant_name': row[2],
                'sample_type': row[3],
                'sampling_date': row[4]
            })

        conn.close()

        return jsonify({
            'total': len(results),
            'results': results
        })

    except Exception as e:
        return jsonify({'error': f'查询失败: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
