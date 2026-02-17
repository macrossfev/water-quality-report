from flask import Blueprint, request, jsonify, session, send_file
from auth import login_required, admin_required, log_operation
from models_v2 import get_db
from datetime import datetime
import json
import os
import sqlite3
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

report_bp = Blueprint('report_bp', __name__)

# ==================== 报告管理 API ====================
@report_bp.route('/api/reports', methods=['GET', 'POST'])
@login_required
def api_reports():
    """报告管理"""
    with get_db() as conn:

        if request.method == 'POST':
            data = request.json
            sample_number = data.get('sample_number')
            company_id = data.get('company_id')
            sample_type_id = data.get('sample_type_id')
            detection_person = data.get('detection_person', '')
            review_person = data.get('review_person', '')
            detection_date = data.get('detection_date')
            remark = data.get('remark', '')
            report_data_list = data.get('data', [])
            template_id = data.get('template_id')
            template_fields = data.get('template_fields', [])
            review_status = data.get('review_status', 'draft')  # 默认为草稿，可以是 'draft' 或 'pending'

            # 新增字段
            report_date = data.get('report_date')
            sample_source = data.get('sample_source', '')
            sampler = data.get('sampler', '')
            sampling_date = data.get('sampling_date')
            sampling_basis = data.get('sampling_basis', '')
            sample_received_date = data.get('sample_received_date')
            sampling_location = data.get('sampling_location', '')
            sample_status = data.get('sample_status', '')
            product_standard = data.get('product_standard', '')
            test_conclusion = data.get('test_conclusion', '')
            detection_items_description = data.get('detection_items_description', '')
            additional_info = data.get('additional_info', '')

            # 获取用户输入的报告编号
            report_number = data.get('report_number', '').strip()

            if not report_number or not sample_number or not sample_type_id:
                return jsonify({'error': '报告编号、样品编号和样品类型不能为空'}), 400

            # 检查报告编号是否已存在
            existing = conn.execute(
                'SELECT id FROM reports WHERE report_number = ?',
                (report_number,)
            ).fetchone()

            if existing:
                return jsonify({'error': f'报告编号 {report_number} 已存在'}), 400

            try:
                cursor = conn.cursor()
                cursor.execute(
                    'INSERT INTO reports (report_number, sample_number, company_id, sample_type_id, '
                    'detection_person, review_person, detection_date, remark, template_id, review_status, created_by, '
                    'report_date, sample_source, sampler, sampling_date, sampling_basis, '
                    'sample_received_date, sampling_location, sample_status, product_standard, '
                    'test_conclusion, detection_items_description, additional_info) '
                    'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                    (report_number, sample_number, company_id, sample_type_id, detection_person,
                     review_person, detection_date, remark, template_id, review_status, session['user_id'],
                     report_date, sample_source, sampler, sampling_date, sampling_basis,
                     sample_received_date, sampling_location, sample_status, product_standard,
                     test_conclusion, detection_items_description, additional_info)
                )
                report_id = cursor.lastrowid

                # 添加报告数据
                for item in report_data_list:
                    if item.get('indicator_id'):
                        cursor.execute(
                            'INSERT INTO report_data (report_id, indicator_id, measured_value, remark) '
                            'VALUES (?, ?, ?, ?)',
                            (report_id, item['indicator_id'], item.get('measured_value', ''),
                             item.get('remark', ''))
                        )

                # 添加模板字段值
                for field in template_fields:
                    if field.get('field_mapping_id') and field.get('field_value'):
                        cursor.execute(
                            'INSERT INTO report_field_values (report_id, field_mapping_id, field_value) '
                            'VALUES (?, ?, ?)',
                            (report_id, field['field_mapping_id'], field['field_value'])
                        )


                status_text = '草稿' if review_status == 'draft' else '提交审核'
                log_operation('创建报告', f'报告编号:{report_number}, 状态:{status_text}', conn=conn)
                return jsonify({'id': report_id, 'report_number': report_number, 'message': '报告创建成功'}), 201
            except Exception as e:
                return jsonify({'error': str(e)}), 500

        # GET请求 - 支持搜索
        search_sample_number = request.args.get('sample_number', '')
        search_company_id = request.args.get('company_id', '')
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))

        query = '''
            SELECT r.*, st.name as sample_type_name, c.name as company_name
            FROM reports r
            LEFT JOIN sample_types st ON r.sample_type_id = st.id
            LEFT JOIN companies c ON r.company_id = c.id
            WHERE 1=1
        '''
        params = []

        if search_sample_number:
            query += ' AND r.sample_number LIKE ?'
            params.append(f'%{search_sample_number}%')

        if search_company_id:
            query += ' AND r.company_id = ?'
            params.append(search_company_id)

        query += ' ORDER BY r.created_at DESC LIMIT ? OFFSET ?'
        params.extend([limit, offset])

        reports = conn.execute(query, params).fetchall()

        return jsonify([dict(report) for report in reports])

@report_bp.route('/api/reports/<int:id>', methods=['GET', 'PUT', 'DELETE'])
@login_required
def api_report_detail(id):
    """报告详情"""
    with get_db() as conn:

        if request.method == 'DELETE':
            # 仅创建人或管理员可删除
            report = conn.execute('SELECT created_by, generated_report_path FROM reports WHERE id = ?', (id,)).fetchone()

            if not report:
                return jsonify({'error': '报告不存在'}), 404

            if session.get('role') not in ('admin', 'super_admin') and report['created_by'] != session['user_id']:
                return jsonify({'error': '无权删除此报告'}), 403

            # 删除生成的报告文件（如果存在）
            if report['generated_report_path'] and os.path.exists(report['generated_report_path']):
                try:
                    os.remove(report['generated_report_path'])
                except Exception as e:
                    print(f"删除报告文件失败: {e}")

            conn.execute('DELETE FROM reports WHERE id = ?', (id,))

            log_operation('删除报告', f'报告ID:{id}', conn=conn)

            return jsonify({'message': '报告删除成功'})

        if request.method == 'PUT':
            # 仅创建人或管理员可修改
            report = conn.execute('SELECT created_by, report_number FROM reports WHERE id = ?', (id,)).fetchone()

            if not report:
                return jsonify({'error': '报告不存在'}), 404

            if session.get('role') not in ('admin', 'super_admin') and report['created_by'] != session['user_id']:
                return jsonify({'error': '无权修改此报告'}), 403

            data = request.json
            sample_number = data.get('sample_number')
            company_id = data.get('company_id')
            detection_person = data.get('detection_person', '')
            review_person = data.get('review_person', '')
            detection_date = data.get('detection_date')
            remark = data.get('remark', '')
            report_data_list = data.get('data', [])
            template_fields = data.get('template_fields', [])

            # 新增字段
            report_date = data.get('report_date')
            sample_source = data.get('sample_source', '')
            sampler = data.get('sampler', '')
            sampling_date = data.get('sampling_date')
            sampling_basis = data.get('sampling_basis', '')
            sample_received_date = data.get('sample_received_date')
            sampling_location = data.get('sampling_location', '')
            sample_status = data.get('sample_status', '')
            product_standard = data.get('product_standard', '')
            test_conclusion = data.get('test_conclusion', '')
            detection_items_description = data.get('detection_items_description', '')
            additional_info = data.get('additional_info', '')

            try:
                cursor = conn.cursor()
                # 更新报告基本信息
                cursor.execute(
                    'UPDATE reports SET sample_number = ?, company_id = ?, detection_person = ?, review_person = ?, '
                    'detection_date = ?, remark = ?, report_date = ?, sample_source = ?, sampler = ?, '
                    'sampling_date = ?, sampling_basis = ?, sample_received_date = ?, sampling_location = ?, '
                    'sample_status = ?, product_standard = ?, test_conclusion = ?, detection_items_description = ?, '
                    'additional_info = ? WHERE id = ?',
                    (sample_number, company_id, detection_person, review_person, detection_date, remark,
                     report_date, sample_source, sampler, sampling_date, sampling_basis,
                     sample_received_date, sampling_location, sample_status, product_standard,
                     test_conclusion, detection_items_description, additional_info, id)
                )

                # 删除旧的报告数据
                cursor.execute('DELETE FROM report_data WHERE report_id = ?', (id,))

                # 插入新的报告数据
                for item in report_data_list:
                    if item.get('indicator_id'):
                        cursor.execute(
                            'INSERT INTO report_data (report_id, indicator_id, measured_value, remark) '
                            'VALUES (?, ?, ?, ?)',
                            (id, item['indicator_id'], item.get('measured_value', ''),
                             item.get('remark', ''))
                        )

                # 删除旧的模板字段值
                cursor.execute('DELETE FROM report_field_values WHERE report_id = ?', (id,))

                # 插入新的模板字段值
                for field in template_fields:
                    if field.get('field_mapping_id') and field.get('field_value'):
                        cursor.execute(
                            'INSERT INTO report_field_values (report_id, field_mapping_id, field_value) '
                            'VALUES (?, ?, ?)',
                            (id, field['field_mapping_id'], field['field_value'])
                        )


                log_operation('更新报告', f'报告编号:{report["report_number"]}', conn=conn)
                return jsonify({'message': '报告更新成功'})
            except Exception as e:
                return jsonify({'error': str(e)}), 500

        # GET请求 - 获取报告详情
        report = conn.execute(
            'SELECT r.*, st.name as sample_type_name, st.code as sample_type_code, '
            'c.name as company_name, u.username as creator_name '
            'FROM reports r '
            'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
            'LEFT JOIN companies c ON r.company_id = c.id '
            'LEFT JOIN users u ON r.created_by = u.id '
            'WHERE r.id = ?',
            (id,)
        ).fetchone()

        if not report:
            return jsonify({'error': '报告不存在'}), 404

        # 获取报告数据
        data = conn.execute(
            'SELECT rd.*, i.name as indicator_name, i.unit, '
            'COALESCE(ti.limit_value, i.limit_value) as limit_value, i.detection_method, '
            'i.group_id, g.name as group_name '
            'FROM report_data rd '
            'LEFT JOIN indicators i ON rd.indicator_id = i.id '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'LEFT JOIN template_indicators ti '
            '    ON ti.indicator_id = rd.indicator_id AND ti.sample_type_id = ? '
            'WHERE rd.report_id = ? '
            'ORDER BY ti.sort_order, g.sort_order, i.sort_order',
            (report['sample_type_id'], id,)
        ).fetchall()

        # 获取模板字段值
        template_fields = []
        if report['template_id']:
            template_fields = conn.execute('''
                SELECT rfv.*, tfm.field_name, tfm.field_display_name
                FROM report_field_values rfv
                LEFT JOIN template_field_mappings tfm ON rfv.field_mapping_id = tfm.id
                WHERE rfv.report_id = ?
            ''', (id,)).fetchall()


        result = dict(report)
        result['data'] = [dict(row) for row in data]
        result['template_fields'] = [dict(row) for row in template_fields]
        return jsonify(result)

# ==================== 模板导入导出 API ====================
@report_bp.route('/api/templates/export', methods=['POST'])
@admin_required
def api_export_template():
    """导出模板JSON"""
    data = request.json
    sample_type_id = data.get('sample_type_id')

    if not sample_type_id:
        return jsonify({'error': '样品类型ID不能为空'}), 400

    with get_db() as conn:

        # 获取样品类型信息
        sample_type = conn.execute(
            'SELECT * FROM sample_types WHERE id = ?',
            (sample_type_id,)
        ).fetchone()

        if not sample_type:
            return jsonify({'error': '样品类型不存在'}), 404

        # 获取关联的检测项目
        template_indicators = conn.execute(
            'SELECT ti.*, i.name as indicator_name, i.unit, i.default_value, i.group_id, '
            'g.name as group_name, COALESCE(ti.limit_value, i.limit_value) as limit_value '
            'FROM template_indicators ti '
            'LEFT JOIN indicators i ON ti.indicator_id = i.id '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'WHERE ti.sample_type_id = ?',
            (sample_type_id,)
        ).fetchall()


        # 构建导出数据
        export_data = {
            'sample_type': dict(sample_type),
            'indicators': [dict(ti) for ti in template_indicators],
            'export_date': datetime.now().isoformat(),
            'version': '2.0'
        }

        # 保存JSON文件
        os.makedirs('exports', exist_ok=True)
        filename = f"exports/template_{sample_type['code']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.json"

        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        log_operation('导出模板', f'导出模板: {sample_type["name"]}')
        return send_file(filename, as_attachment=True, download_name=f"template_{sample_type['code']}.json")

@report_bp.route('/api/templates/import', methods=['POST'])
@admin_required
def api_import_template():
    """导入模板JSON"""
    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    try:
        content = file.read().decode('utf-8')
        data = json.loads(content)

        sample_type_data = data.get('sample_type')
        indicators_data = data.get('indicators', [])

        if not sample_type_data:
            return jsonify({'error': 'JSON格式错误:缺少sample_type'}), 400

        with get_db() as conn:
            cursor = conn.cursor()

            # 检查样品类型是否已存在
            existing = cursor.execute(
                'SELECT id FROM sample_types WHERE code = ?',
                (sample_type_data['code'],)
            ).fetchone()

            if existing:
                return jsonify({'error': f'样品类型代码 {sample_type_data["code"]} 已存在'}), 400

            # 创建样品类型
            cursor.execute(
                'INSERT INTO sample_types (name, code, description) VALUES (?, ?, ?)',
                (sample_type_data['name'], sample_type_data['code'], sample_type_data.get('description', ''))
            )
            sample_type_id = cursor.lastrowid

            # 导入检测项目(需要匹配现有的indicator)
            imported_count = 0
            for item in indicators_data:
                # 优先按(name, group_id)精确匹配，回退到按name匹配
                indicator = None
                if item.get('group_id'):
                    indicator = cursor.execute(
                        'SELECT id FROM indicators WHERE name = ? AND group_id = ?',
                        (item['indicator_name'], item['group_id'])
                    ).fetchone()
                if not indicator:
                    indicator = cursor.execute(
                        'SELECT id FROM indicators WHERE name = ? LIMIT 1',
                        (item['indicator_name'],)
                    ).fetchone()

                if indicator:
                    try:
                        cursor.execute(
                            'INSERT INTO template_indicators (sample_type_id, indicator_id, is_required, sort_order, limit_value) '
                            'VALUES (?, ?, ?, ?, ?)',
                            (sample_type_id, indicator['id'], item.get('is_required', False),
                             item.get('sort_order', 0), item.get('limit_value', ''))
                        )
                        imported_count += 1
                    except sqlite3.IntegrityError:
                        pass  # 忽略重复项


            log_operation('导入模板', f'导入模板: {sample_type_data["name"]}, 检测项:{imported_count}', conn=conn)
            return jsonify({
                'message': f'模板导入成功,共导入 {imported_count} 个检测项目',
                'sample_type_id': sample_type_id
            })

    except json.JSONDecodeError:
        return jsonify({'error': 'JSON格式错误'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== 报告导出 API ====================
@report_bp.route('/api/reports/<int:id>/export/excel', methods=['GET'])
@login_required
def api_export_excel(id):
    """导出Excel报告"""
    with get_db() as conn:

        report = conn.execute(
            'SELECT r.*, st.name as sample_type_name, c.name as company_name '
            'FROM reports r '
            'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
            'LEFT JOIN companies c ON r.company_id = c.id '
            'WHERE r.id = ?',
            (id,)
        ).fetchone()

        if not report:
            return jsonify({'error': '报告不存在'}), 404

        data = conn.execute(
            'SELECT rd.*, i.name as indicator_name, i.unit, g.name as group_name '
            'FROM report_data rd '
            'LEFT JOIN indicators i ON rd.indicator_id = i.id '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'LEFT JOIN template_indicators ti '
            '    ON ti.indicator_id = rd.indicator_id AND ti.sample_type_id = ? '
            'WHERE rd.report_id = ? '
            'ORDER BY ti.sort_order, g.sort_order, i.sort_order',
            (report['sample_type_id'], id,)
        ).fetchall()


        # 创建Excel工作簿
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "水质检测报告"

        # 设置样式
        title_font = Font(name='宋体', size=16, bold=True)
        header_font = Font(name='宋体', size=11, bold=True)
        normal_font = Font(name='宋体', size=10)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 标题
        ws.merge_cells('A1:G1')
        title_cell = ws['A1']
        title_cell.value = '水质检测报告'
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30

        # 报告信息
        row = 3
        info_items = [
            ('报告编号', report['report_number']),
            ('样品编号', report['sample_number']),
            ('样品类型', report['sample_type_name']),
            ('委托单位', report['company_name']),
            ('检测日期', report['detection_date']),
            ('检测人员', report['detection_person']),
            ('审核人员', report['review_person'])
        ]

        for label, value in info_items:
            if value:
                ws[f'A{row}'] = label + '：'
                ws[f'B{row}'] = value
                ws[f'A{row}'].font = header_font
                ws.merge_cells(f'B{row}:G{row}')
                row += 1

        row += 1

        # 表头
        headers = ['序号', '检测项目', '单位', '检测结果', '所属分组', '备注']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=row, column=col)
            cell.value = header
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

        # 数据行
        for idx, item in enumerate(data, start=1):
            row += 1
            row_data = [
                idx,
                item['indicator_name'],
                item['unit'] or '',
                item['measured_value'] or '',
                item['group_name'] or '',
                item['remark'] or ''
            ]

            for col, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row, column=col)
                cell.value = value
                cell.font = normal_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = border

        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 20

        # 保存文件
        os.makedirs('exports', exist_ok=True)
        filename = f"exports/report_{report['report_number']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
        wb.save(filename)

        log_operation('导出Excel报告', f'报告编号:{report["report_number"]}')
        return send_file(filename, as_attachment=True, download_name=f"{report['report_number']}.xlsx")

@report_bp.route('/api/reports/<int:id>/export/word', methods=['GET'])
@login_required
def api_export_word(id):
    """导出Word报告"""
    with get_db() as conn:

        report = conn.execute(
            'SELECT r.*, st.name as sample_type_name, c.name as company_name '
            'FROM reports r '
            'LEFT JOIN sample_types st ON r.sample_type_id = st.id '
            'LEFT JOIN companies c ON r.company_id = c.id '
            'WHERE r.id = ?',
            (id,)
        ).fetchone()

        if not report:
            return jsonify({'error': '报告不存在'}), 404

        data = conn.execute(
            'SELECT rd.*, i.name as indicator_name, i.unit, g.name as group_name '
            'FROM report_data rd '
            'LEFT JOIN indicators i ON rd.indicator_id = i.id '
            'LEFT JOIN indicator_groups g ON i.group_id = g.id '
            'LEFT JOIN template_indicators ti '
            '    ON ti.indicator_id = rd.indicator_id AND ti.sample_type_id = ? '
            'WHERE rd.report_id = ? '
            'ORDER BY ti.sort_order, g.sort_order, i.sort_order',
            (report['sample_type_id'], id,)
        ).fetchall()


        # 创建Word文档
        doc = Document()

        # 标题
        title = doc.add_heading('水质检测报告', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告信息
        info_items = [
            ('报告编号', report['report_number']),
            ('样品编号', report['sample_number']),
            ('样品类型', report['sample_type_name']),
            ('委托单位', report['company_name'] or ''),
            ('检测日期', report['detection_date'] or ''),
            ('检测人员', report['detection_person'] or ''),
            ('审核人员', report['review_person'] or '')
        ]

        for label, value in info_items:
            if value:
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)

        doc.add_paragraph()

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ['序号', '检测项目', '单位', '检测结果', '所属分组', '备注']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 数据行
        for idx, item in enumerate(data, start=1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item['indicator_name']
            row.cells[2].text = item['unit'] or ''
            row.cells[3].text = item['measured_value'] or ''
            row.cells[4].text = item['group_name'] or ''
            row.cells[5].text = item['remark'] or ''

            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文件
        os.makedirs('exports', exist_ok=True)
        filename = f"exports/report_{report['report_number']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(filename)

        log_operation('导出Word报告', f'报告编号:{report["report_number"]}')
        return send_file(filename, as_attachment=True, download_name=f"{report['report_number']}.docx")
